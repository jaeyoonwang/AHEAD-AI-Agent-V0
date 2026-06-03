#!/usr/bin/env python3
"""Generates AHEAD_AI_UX_Workflow.docx — UX and workflow specification.
Run: pip install python-docx && python3 generate_ux_doc.py
"""
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = pathlib.Path(__file__).parent

doc = Document()
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

def h1(t):
    p = doc.add_heading(t, 1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
    p.runs[0].font.size = Pt(15)
def h2(t):
    p = doc.add_heading(t, 2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
    p.runs[0].font.size = Pt(12)
def body(t):
    p = doc.add_paragraph(t)
    if p.runs: p.runs[0].font.size = Pt(10.5)
def bullet(t, lv=0):
    p = doc.add_paragraph(t, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (lv + 1))
def mono(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(t)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shd)
def tbl(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].text = h
        run = hc[i].paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'D6E4F0')
        hc[i]._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            if cells[ci].paragraphs[0].runs:
                cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
def sp(): doc.add_paragraph('')


# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('AHEAD AI Agent — UX & Workflow Specification', 0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
title.runs[0].font.size = Pt(18)
sub = doc.add_paragraph('SMS Notification Design, Escalation Logic, and Issue Log  |  UNICEF AHEAD × Gates Foundation AI Fellows  |  June 2026')
sub.runs[0].font.size = Pt(9); sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88); sub.runs[0].italic = True
sp()

# ── 1. User Roles ─────────────────────────────────────────────────────────────
h1('1. User Roles and Notification Scope')
body('Six roles exist in the AHEAD hierarchy. For MVP, only facility workers and woreda HMIS officers receive active SMS notifications. Zone, regional, and national staff receive escalation summaries only when issues remain unresolved at lower levels. The AHEAD/UNICEF team monitors everything passively via the issue log web page.')
sp()
tbl(
    ['Role', 'Notification channel', 'What they receive', 'MVP scope'],
    [
        ['Facility worker',
         'SMS (individual)',
         'One SMS per issue with a reference ID; confirmation loop; resolution confirmation',
         'Active — first point of contact for every issue'],
        ['Woreda HMIS officer',
         'SMS (summary)',
         'Batched summary of all unresolved issues in their woreda after 72h',
         'Active — escalation target'],
        ['Zone officer',
         'SMS (summary)',
         'Summary of unresolved issues across woredas after 1 week',
         'Active — escalation target'],
        ['Regional focal point',
         'SMS (summary)',
         'Summary of unresolved issues across zones after 1 further week',
         'Active — escalation target'],
        ['National TWG',
         'SMS (monthly digest)',
         'End-of-month summary of all unresolved issues across Ethiopia',
         'Active — monthly digest only'],
        ['UNICEF / AHEAD team',
         'Issue log web page',
         'Real-time view of all issues, statuses, and full SMS threads',
         'Passive — observe and intervene via log page'],
    ]
)
sp()
body('Important distinction: the notification contacts (phone numbers) are stored in the agent\'s own contact registry, not in DHIS2 user accounts. DHIS2 user accounts control data-entry access. The contact registry controls who gets notified. These are maintained separately.')
sp()
body('Demo focus: The live demonstration centres entirely on the facility-level health worker — the person who receives the initial SMS, engages in the correction conversation, and sees their submission corrected in DHIS2. The escalation cascade (woreda → zone → region → national) is part of the production architecture and fires automatically in the background, but is not walked through in the demo. The AHEAD team monitors everything via the standalone issue log web app shown at the end of the demo.')
sp()

# ── 2. Notification Triggers ──────────────────────────────────────────────────
h1('2. Notification Triggers')
body('The MVP implements two of the five AHEAD detection methods: SD-based outlier detection (Method 1, AHEAD Table 1) and absolute difference checks (Method 5). DQ thresholds are aligned with AHEAD operational guidance: outlier detection fires when a value exceeds 2.0 standard deviations from the facility historical baseline OR differs by more than 100 doses in absolute terms. DTP1/DTP3 inconsistency fires when Penta3 exceeds Penta1 by more than 30% relative OR more than 100 doses absolute.')
sp()
tbl(
    ['Issue type', 'What causes it', 'Detection threshold', 'Detection timing', 'Who is notified first'],
    [
        ['Statistical outlier',
         'A submitted value deviates from the facility historical baseline',
         'Z-score > 2.0 SD (AHEAD Table 1 operational threshold) OR absolute diff > 100 doses',
         'Within ~30 seconds of form submission',
         'Facility worker'],
        ['DTP1/DTP3 inconsistency',
         'Penta3 doses exceed Penta1 doses for same facility-period (biologically implausible)',
         '> 30% relative gap OR > 100 absolute dose gap (monthly facility threshold)',
         'Within ~30 seconds of form submission',
         'Facility worker'],
        ['Missing report',
         'Facility has no completed dataset registration by day 5 of the following month',
         'No completeDataSetRegistration record in DHIS2 for the period',
         'Daily at 8am, starting day 5 of following month',
         'Facility worker'],
    ]
)
sp()

# ── 3. Message Templates ──────────────────────────────────────────────────────
h1('3. SMS Message Templates')
body('All outbound notifications present numbered multiple-choice options — the same response categories used in the AHEAD reference guide (sections 2.2–2.4). The worker replies with a number (1, 2, 3 etc.). Options that modify DHIS2 data require a YES/NO confirmation before the write-back is applied. Options requiring additional input (a specific value or a brief explanation) trigger a short follow-up question before the confirmation step.')
sp()

h2('3.1  Outlier — initial notification (6 options)')
body('Sent when a data value is flagged by outlier detection. Aligned with AHEAD guide section 2.2. Options 1, 3, and 4 write data back to DHIS2 and require a YES/NO confirmation. Options 5 and 6 are noted for HQ manual adjustment with no auto write-back. Option 2 resolves immediately with no data change.')
mono(
    'AHEAD DQ Alert [DQ-4XDF]\n'
    'Addi Arekay Health Center — Jun 2026\n'
    'BCG: 970 doses (expected 79–114)\n'
    'Reply with option number:\n'
    '1. Replace with 6-month average\n'
    '2. Keep as-is (no action)\n'
    '3. Set to zero\n'
    '4. Replace with specific value\n'
    '5. At health facility doses only\n'
    '6. Outreach doses only'
)
sp()

h2('3.2  Outlier — follow-up for option 4 (specific value)')
body('Agent asks for the corrected value before proceeding to the YES/NO confirmation step.')
mono(
    '[DQ-4XDF] Please reply with the correct value\n'
    '(numbers only, e.g. 97).'
)
sp()

h2('3.3  Outlier — YES/NO confirmation before DHIS2 write-back')
body('Shown before any data modification. The confirmation message states exactly what will change. If the user replies NO, the original alert is re-sent and they can choose a different option.')
mono(
    '[DQ-4XDF] Confirm change:\n'
    'BCG <1yr — Addi Arekay Health Center (Jun 2026)\n'
    '970 -> 97\n'
    'Reply YES to update DHIS2 or NO to re-enter.'
)
sp()
body('For option 1 (6-month average), the agent auto-computes the average of 3 periods before and 3 after the flagged period and shows the computed value in the confirmation:')
mono(
    '[DQ-4XDF] Confirm: Replace with 6-month average (96)\n'
    'Addi Arekay Health Center (Jun 2026)\n'
    'Reply YES to update DHIS2 or NO to choose again.'
)
sp()
body('For option 3 (set to zero), the confirmation shows the explicit 0 value:')
mono(
    '[DQ-4XDF] Confirm: Set BCG <1yr to 0\n'
    'Addi Arekay Health Center (Jun 2026)\n'
    'Reply YES to update DHIS2 or NO to choose again.'
)
sp()

h2('3.4  Outlier — resolution confirmation')
body('Sent after the user replies YES and the write-back is applied.')
mono(
    '[DQ-4XDF] Noted. Value corrected to 97 in DHIS2. Thank you.'
)
sp()
body('For option 1 (6-month average):')
mono(
    '[DQ-4XDF] Noted. Replaced with 6-month average (96) in DHIS2. Thank you.'
)
sp()
body('For options 5 and 6 (no auto write-back), no confirmation step is triggered — the agent closes with a note:')
mono(
    '[DQ-4XDF] Noted. Flagged for HQ manual adjustment\n'
    '(at health facility doses only). Issue logged. Thank you.'
)
sp()

h2('3.5  DTP1/DTP3 inconsistency — initial notification (5 options)')
body('Sent when Penta3 > Penta1. Aligned with AHEAD guide section 2.3. Options 2 and 3 set DTP3 = DTP1 or DTP1 = DTP3 respectively; option 4 uses a user-supplied value. All three require YES/NO confirmation. Option 1 resolves immediately. Option 5 is noted with no auto write-back.')
mono(
    'AHEAD DQ Alert [DQ-BK19]\n'
    'Addi Arekay Health Center — Jun 2026\n'
    'Penta1 under-1: 60 | Penta3 under-1: 90\n'
    'Penta3 exceeds Penta1 by 50%.\n'
    'Reply with option number:\n'
    '1. Keep as-is\n'
    '2. Use DTP1 value for both\n'
    '3. Use DTP3 value for both\n'
    '4. Replace with specific value\n'
    '5. Other'
)
sp()

h2('3.6  DTP — follow-up for option 4 (specific value)')
mono(
    '[DQ-BK19] Please reply with the correct value\n'
    '(numbers only, e.g. 72).'
)
sp()

h2('3.7  DTP — YES/NO confirmation before DHIS2 write-back')
body('Examples for options 2, 3, and 4 respectively:')
mono(
    '[DQ-BK19] Confirm change:\n'
    'Penta3 <1yr — Addi Arekay Health Center (Jun 2026)\n'
    '90 -> 60 (set to DTP1 value)\n'
    'Reply YES to update DHIS2 or NO to choose again.'
)
sp()
mono(
    '[DQ-BK19] Confirm change:\n'
    'Penta1 <1yr — Addi Arekay Health Center (Jun 2026)\n'
    '60 -> 90 (set to DTP3 value)\n'
    'Reply YES to update DHIS2 or NO to choose again.'
)
sp()
mono(
    '[DQ-BK19] Confirm change:\n'
    'Penta3 <1yr — Addi Arekay Health Center (Jun 2026)\n'
    '90 -> 72\n'
    'Reply YES to update DHIS2 or NO to re-enter.'
)
sp()

h2('3.8  DTP — resolution confirmation')
mono(
    '[DQ-BK19] Noted. Value corrected to 60 in DHIS2. Thank you.'
)
sp()
body('For option 5 (no auto write-back):')
mono(
    '[DQ-BK19] Noted. Flagged for HQ review. Issue logged. Thank you.'
)
sp()

h2('3.9  Missing report — initial notification (SUBMIT + 4 options)')
body('Aligned with AHEAD guide section 2.4. Recovery before imputation: the SUBMIT option is presented first so the worker can locate and submit the missing report. Cleaning options are a fallback only if recovery fails. Options 1, 2, 3, and 4 do not write data automatically — options 2 and 3 are flagged for HQ imputation; option 1 is noted; option 4 is noted. SUBMIT acknowledges and closes the issue once the submission appears.')
mono(
    'AHEAD DQ Alert [DQ-MN33]\n'
    'Addi Arekay Health Center\n'
    'No completed EPI report for May 2026.\n'
    'Can you submit the report? Reply SUBMIT.\n'
    'Or if data is unavailable, select:\n'
    '1. Will submit by [date]\n'
    '2. Data cannot be recovered\n'
    '3. Facility closed / no service that month\n'
    '4. Other'
)
sp()

h2('3.10  Missing report — follow-up for SUBMIT')
mono(
    '[DQ-MN33] Thanks — checking DHIS2 within the hour.\n'
    'Issue will close automatically once your\n'
    'submission appears.'
)
sp()

h2('3.11  Missing report — follow-up for option 1 (will submit by date)')
body('Agent asks for the target date before acknowledging.')
mono(
    '[DQ-MN33] Please reply with the date you will submit\n'
    '(e.g., 2026-06-10).'
)
sp()
body('After the user provides the date, agent acknowledges and logs:')
mono(
    '[DQ-MN33] Noted. Submission expected by 2026-06-10.\n'
    'Issue logged for follow-up. Thank you.'
)
sp()

h2('3.12  Missing report — options 2, 3, and 4 (no auto write-back)')
body('Options 2 and 3 are flagged for HQ zero imputation. Option 4 is noted. None trigger a DHIS2 write-back.')
mono(
    '[DQ-MN33] Noted. Data cannot be recovered — flagged\n'
    'for HQ imputation. Issue logged. Thank you.'
)
sp()
mono(
    '[DQ-MN33] Noted. Facility closed / no service that month\n'
    '— flagged for HQ zero imputation. Issue logged. Thank you.'
)
sp()
mono(
    '[DQ-MN33] Noted. Issue logged for HQ review. Thank you.'
)
sp()

h2('3.13  Retry (24h no reply — up to 3 times)')
mono(
    '[DQ-4XDF] Reminder (2 of 3):\n'
    'Addi Arekay HC, Jun 2026 — BCG under-1 = 970\n'
    'Reply 1-6 to resolve. Escalates if no reply.'
)
sp()

h2('3.14  Woreda escalation summary')
body('Sent to the woreda HMIS officer after 72h of no facility-level resolution. One SMS per woreda covers all open issues in that woreda.')
mono(
    '[DQ-WOREDA] Addi Arekay Woreda\n'
    '2 unresolved issues, Jun 2026:\n'
    'DQ-4XDF Addi Arekay HC: BCG outlier (3 days)\n'
    'DQ-BK19 Addi Arekay HC: Penta3 > Penta1 (3 days)\n'
    'Details: <issue log URL>'
)
sp()

h2('3.15  Monthly national digest')
body('Sent to the national TWG on the first of each month covering all prior-month unresolved issues.')
mono(
    '[DQ-MONTHLY] Jun 2026 Summary\n'
    'Ethiopia: 2 issues still open\n'
    'Amhara / N. Gondar:\n'
    '  - 1 outlier (Addi Arekay HC, 25 days)\n'
    '  - 1 DTP inconsistency (Addi Arekay HC)\n'
    'View all: <issue log URL>'
)
sp()

# ── 4. Conversation Flows ─────────────────────────────────────────────────────
h1('4. Conversation Flows')
body('The numbered-option model keeps conversations short but includes a mandatory YES/NO confirmation before any DHIS2 write-back. Options that modify data take 3–4 turns: notification → option selection → (optional follow-up for extra input) → YES/NO confirmation → resolution. Options that do not modify data (keep as-is, flagged for HQ) resolve in 2 turns. If the user replies NO at the confirmation step, the original alert is re-sent and they may choose a different option.')
sp()

h2('4.1  Outlier — option 1: 6-month average (3 turns)')
body('Agent auto-computes the average of 3 periods before and 3 after the flagged period, presents it in the confirmation, and applies it on YES. No user value input needed.')
mono(
    'Agent     AHEAD DQ Alert [DQ-4XDF]\n'
    '          Addi Arekay Health Center — Jun 2026\n'
    '          BCG: 970 doses (expected 79–114)\n'
    '          Reply with option number:\n'
    '          1. Replace with 6-month average\n'
    '          2. Keep as-is (no action)\n'
    '          3. Set to zero\n'
    '          4. Replace with specific value\n'
    '          5. At health facility doses only\n'
    '          6. Outreach doses only\n'
    '\n'
    'Facility  1\n'
    '\n'
    'Agent     [DQ-4XDF] Confirm: Replace with 6-month average (96)\n'
    '          Addi Arekay Health Center (Jun 2026)\n'
    '          Reply YES to update DHIS2 or NO to choose again.\n'
    '\n'
    'Facility  YES\n'
    '\n'
    'Agent     [DQ-4XDF] Noted. Replaced with 6-month average (96) in DHIS2. Thank you.'
)
sp()

h2('4.2  Outlier — option 2: keep as-is (2 turns — no data change)')
body('Resolves immediately with no DHIS2 modification and no confirmation step.')
mono(
    'Agent     AHEAD DQ Alert [DQ-4XDF]\n'
    '          Addi Arekay Health Center — Jun 2026\n'
    '          BCG: 970 doses (expected 79–114)\n'
    '          Reply with option number:\n'
    '          1. Replace with 6-month average\n'
    '          2. Keep as-is (no action)\n'
    '          ...\n'
    '\n'
    'Facility  2\n'
    '\n'
    'Agent     [DQ-4XDF] Noted. No change applied. Issue resolved. Thank you.'
)
sp()

h2('4.3  Outlier — option 4: replace with specific value (4 turns)')
body('Four turns: notification, option selection, value entry, YES/NO confirmation. This is the full flow for any option that requires user-supplied input before DHIS2 write-back.')
mono(
    'Agent     AHEAD DQ Alert [DQ-4XDF]\n'
    '          Addi Arekay Health Center — Jun 2026\n'
    '          BCG: 970 doses (expected 79–114)\n'
    '          Reply with option number:\n'
    '          1. Replace with 6-month average\n'
    '          2. Keep as-is (no action)\n'
    '          3. Set to zero\n'
    '          4. Replace with specific value\n'
    '          5. At health facility doses only\n'
    '          6. Outreach doses only\n'
    '\n'
    'Facility  4\n'
    '\n'
    'Agent     [DQ-4XDF] Please reply with the correct value\n'
    '          (numbers only, e.g. 97).\n'
    '\n'
    'Facility  97\n'
    '\n'
    'Agent     [DQ-4XDF] Confirm change:\n'
    '          BCG <1yr — Addi Arekay Health Center (Jun 2026)\n'
    '          970 -> 97\n'
    '          Reply YES to update DHIS2 or NO to re-enter.\n'
    '\n'
    'Facility  YES\n'
    '\n'
    'Agent     [DQ-4XDF] Noted. Value corrected to 97 in DHIS2. Thank you.'
)
sp()

h2('4.4  Outlier — option 4, user replies NO at confirmation (re-enter flow)')
body('If the user replies NO at the confirmation step, the original alert is re-sent and they can choose a different option.')
mono(
    'Facility  NO\n'
    '\n'
    'Agent     [DQ-4XDF] No change applied. Re-sending alert:\n'
    '          AHEAD DQ Alert [DQ-4XDF]\n'
    '          Addi Arekay Health Center — Jun 2026\n'
    '          BCG: 970 doses (expected 79–114)\n'
    '          Reply with option number:\n'
    '          1. Replace with 6-month average\n'
    '          2. Keep as-is (no action)\n'
    '          3. Set to zero\n'
    '          4. Replace with specific value\n'
    '          5. At health facility doses only\n'
    '          6. Outreach doses only'
)
sp()

h2('4.5  Outlier — options 5 and 6 (noted for HQ, no write-back, 2 turns)')
body('These options do not trigger a confirmation step — the agent logs the decision and closes.')
mono(
    'Facility  5\n'
    '\n'
    'Agent     [DQ-4XDF] Noted. Flagged for HQ manual adjustment\n'
    '          (at health facility doses only). Issue logged. Thank you.'
)
sp()

h2('4.6  DTP inconsistency — option 2: use DTP1 value for both (3 turns)')
body('Agent sets DTP3 = DTP1 value. Requires YES/NO confirmation before writing.')
mono(
    'Agent     AHEAD DQ Alert [DQ-BK19]\n'
    '          Addi Arekay Health Center — Jun 2026\n'
    '          Penta1 under-1: 60 | Penta3 under-1: 90\n'
    '          Penta3 exceeds Penta1 by 50%.\n'
    '          Reply with option number:\n'
    '          1. Keep as-is\n'
    '          2. Use DTP1 value for both\n'
    '          3. Use DTP3 value for both\n'
    '          4. Replace with specific value\n'
    '          5. Other\n'
    '\n'
    'Facility  2\n'
    '\n'
    'Agent     [DQ-BK19] Confirm change:\n'
    '          Penta3 <1yr — Addi Arekay Health Center (Jun 2026)\n'
    '          90 -> 60 (set to DTP1 value)\n'
    '          Reply YES to update DHIS2 or NO to choose again.\n'
    '\n'
    'Facility  YES\n'
    '\n'
    'Agent     [DQ-BK19] Noted. Value corrected to 60 in DHIS2. Thank you.'
)
sp()

h2('4.7  DTP inconsistency — option 4: replace with specific value (4 turns)')
body('Four turns: notification, option selection, value entry, YES/NO confirmation.')
mono(
    'Agent     AHEAD DQ Alert [DQ-BK19]\n'
    '          Addi Arekay Health Center — Jun 2026\n'
    '          Penta1 under-1: 60 | Penta3 under-1: 90\n'
    '          Penta3 exceeds Penta1 by 50%.\n'
    '          Reply with option number:\n'
    '          1. Keep as-is\n'
    '          2. Use DTP1 value for both\n'
    '          3. Use DTP3 value for both\n'
    '          4. Replace with specific value\n'
    '          5. Other\n'
    '\n'
    'Facility  4\n'
    '\n'
    'Agent     [DQ-BK19] Please reply with the correct value\n'
    '          (numbers only, e.g. 72).\n'
    '\n'
    'Facility  72\n'
    '\n'
    'Agent     [DQ-BK19] Confirm change:\n'
    '          Penta3 <1yr — Addi Arekay Health Center (Jun 2026)\n'
    '          90 -> 72\n'
    '          Reply YES to update DHIS2 or NO to re-enter.\n'
    '\n'
    'Facility  YES\n'
    '\n'
    'Agent     [DQ-BK19] Noted. Value corrected to 72 in DHIS2. Thank you.'
)
sp()

h2('4.8  DTP inconsistency — option 1: keep as-is (2 turns — no data change)')
body('Resolves immediately with no DHIS2 modification.')
mono(
    'Facility  1\n'
    '\n'
    'Agent     [DQ-BK19] Noted. No change applied. Issue resolved. Thank you.'
)
sp()

h2('4.9  DTP inconsistency — option 5: other (noted, no write-back, 2 turns)')
mono(
    'Facility  5\n'
    '\n'
    'Agent     [DQ-BK19] Noted. Flagged for HQ review. Issue logged. Thank you.'
)
sp()

h2('4.10  Missing report — recovery succeeds (SUBMIT)')
body('Worker submits the missing report in DHIS2. Agent re-checks within the hour and closes automatically.')
mono(
    'Agent     AHEAD DQ Alert [DQ-MN33]\n'
    '          Addi Arekay Health Center\n'
    '          No completed EPI report for May 2026.\n'
    '          Can you submit the report? Reply SUBMIT.\n'
    '          Or if data is unavailable, select:\n'
    '          1. Will submit by [date]\n'
    '          2. Data cannot be recovered\n'
    '          3. Facility closed / no service that month\n'
    '          4. Other\n'
    '\n'
    'Facility  SUBMIT\n'
    '\n'
    'Agent     [DQ-MN33] Thanks — checking DHIS2 within the hour.\n'
    '          Issue will close automatically once your\n'
    '          submission appears.\n'
    '\n'
    '          [1 hour later — agent finds the registration]\n'
    '\n'
    'Agent     [DQ-MN33] Confirmed. Addi Arekay HC May 2026 report received. Issue closed.'
)
sp()

h2('4.11  Missing report — option 1: will submit by date (3 turns)')
body('User provides a commitment date. Agent acknowledges and flags for follow-up. No DHIS2 write-back.')
mono(
    'Facility  1\n'
    '\n'
    'Agent     [DQ-MN33] Please reply with the date you will submit\n'
    '          (e.g., 2026-06-10).\n'
    '\n'
    'Facility  2026-06-10\n'
    '\n'
    'Agent     [DQ-MN33] Noted. Submission expected by 2026-06-10.\n'
    '          Issue logged for follow-up. Thank you.'
)
sp()

h2('4.12  Missing report — options 2, 3, 4 (noted, flagged for HQ, 2 turns)')
body('None of these options trigger a DHIS2 write-back. Options 2 and 3 are flagged for HQ zero imputation; option 4 is noted for HQ review.')
mono(
    'Facility  3\n'
    '\n'
    'Agent     [DQ-MN33] Noted. Facility closed / no service that month\n'
    '          — flagged for HQ zero imputation. Issue logged. Thank you.'
)
sp()

h2('4.13  No reply — retry then escalation')
mono(
    'Day 0     Agent sends numbered-option notification\n'
    'Day 1     No reply — Retry 1: "[DQ-4XDF] Reminder (1 of 3): Addi Arekay HC BCG=970. Reply 1-6."\n'
    'Day 2     No reply — Retry 2: "[DQ-4XDF] Reminder (2 of 3): ... Issue escalates if no reply."\n'
    'Day 3     No reply — Retry 3 + escalate: woreda HMIS officer receives summary SMS\n'
    'Day 10    No woreda resolution — escalate to zone\n'
    'Day 17    No zone resolution — escalate to region\n'
    'Month end National digest includes all unresolved issues'
)
sp()

# ── 5. Conversation States ─────────────────────────────────────────────────────
h1('5. Conversation States')
body('Each issue progresses through four internal conversation states. The state determines how the agent interprets an inbound reply from the facility worker.')
sp()
tbl(
    ['State', 'Meaning', 'Agent expects', 'Next state on valid input'],
    [
        ['awaiting_option',
         'Alert has been sent; waiting for the worker to pick a numbered option',
         'A digit matching one of the presented options',
         'awaiting_followup (if option needs input) or awaiting_confirmation (if option writes data directly) or closed (if no-write option)'],
        ['awaiting_followup',
         'Worker picked an option that requires additional input (e.g. a specific value or a date)',
         'The required input (a number, a date, etc.)',
         'awaiting_confirmation (if option writes data) or closed (if HQ-flagged option)'],
        ['awaiting_confirmation',
         'Agent has shown the exact proposed change; waiting for YES/NO before writing to DHIS2',
         'YES or NO. If NO, issue reverts to awaiting_option and the alert is re-sent.',
         'closed (on YES) or awaiting_option (on NO)'],
        ['closed',
         'Issue is fully resolved; no further action expected',
         '(no further input expected — any reply is ignored or acknowledged)',
         '—'],
    ]
)
sp()
body('State transition summary for a 4-turn option (e.g., outlier option 4): awaiting_option → awaiting_followup → awaiting_confirmation → closed. For a 3-turn option (e.g., outlier option 1): awaiting_option → awaiting_confirmation → closed. For a 2-turn option (e.g., keep as-is): awaiting_option → closed.')
sp()

# ── 6. Issue Log Status Progression ──────────────────────────────────────────
h1('6. Issue Log Status Progression')
body('Issue statuses visible on the dashboard at http://localhost:5001/issues. The dashboard auto-refreshes every 10 seconds.')
sp()
tbl(
    ['Status', 'Meaning'],
    [
        ['OPEN',        'Issue detected; no notification sent yet (brief window between detection and SMS delivery)'],
        ['NOTIFIED',    'Initial alert SMS sent to facility worker; waiting for first reply'],
        ['IN PROGRESS', 'First reply received from facility worker; conversation is underway'],
        ['CONFIRMING',  'Worker has provided a value or selected a write-back option; waiting for YES/NO confirmation before DHIS2 write-back'],
        ['RESOLVED',    'Issue closed after YES confirmation and successful DHIS2 write-back'],
        ['CONFIRMED OK', 'Issue closed with no data change (keep as-is or equivalent option)'],
        ['ESCALATED',   'No facility response after 3 retries; woreda (or higher) officer notified'],
    ]
)
sp()
body('Status progression for a typical resolved issue: OPEN → NOTIFIED → IN PROGRESS → CONFIRMING → RESOLVED. For a keep-as-is option: OPEN → NOTIFIED → IN PROGRESS → CONFIRMED OK. For a no-response escalation: OPEN → NOTIFIED → ESCALATED.')
sp()

# ── 7. Escalation Timeline ────────────────────────────────────────────────────
h1('7. Escalation Timeline')
body('The timeline below shows a single issue from detection to national digest, assuming no resolution at any level.')
sp()
mono(
    'Day 0    Issue detected → facility worker notified (SMS)\n'
    '         ├── Facility replies → confirmation loop → resolved (ideal path)\n'
    '         └── No reply...\n'
    '\n'
    'Day 1    Retry 1 of 3 (24h after initial)\n'
    'Day 2    Retry 2 of 3 (48h after initial)\n'
    'Day 3    Retry 3 of 3 + escalate to woreda HMIS officer\n'
    '         ├── Woreda intervenes → resolved\n'
    '         └── No resolution...\n'
    '\n'
    'Day 10   Escalate to zone officer (1 week after woreda notification)\n'
    '         ├── Zone intervenes → resolved\n'
    '         └── No resolution...\n'
    '\n'
    'Day 17   Escalate to regional focal point (1 week after zone notification)\n'
    '         └── No resolution...\n'
    '\n'
    'Month end  National TWG receives monthly digest (includes this issue)'
)
sp()
body('Escalation messages at woreda level and above are sent as a single summary SMS covering all unresolved issues from that org unit in the same period — not one SMS per issue. This prevents notification fatigue at higher levels where one officer may receive reports from many facilities.')
sp()

# ── 8. Issue Log Web App ──────────────────────────────────────────────────────
h1('8. Issue Log Web App (Standalone)')
body('A standalone web application at http://localhost:5001/issues, completely separate from DHIS2. No DHIS2 login required. It reads from the agent\'s own SQLite database and is the primary monitoring surface for the AHEAD team and supervisors. Every issue the agent detects, every SMS sent and received, and every state change is visible here in real-time. The dashboard auto-refreshes every 10 seconds.')
sp()

h2('Page layout (wireframe)')
mono(
    'AHEAD Data Issue Log                      [ Open: 2 | Resolved: 47 | All: 49 ]\n'
    'Filter: [ All types v ]  [ All woredas v ]  [ Apr 2026 v ]  [ Search ref ID... ]\n'
    '--------------------------------------------------------------------------------\n'
    'Ref     Facility          Issue              Period  Status    Opened      Updated\n'
    '--------------------------------------------------------------------------------\n'
    'DQ-4XDF Addi Arekay HC    Outlier: BCG 970   202606  RESOLVED  2026-06-01  06-01\n'
    'DQ-BK19 Addi Arekay HC    DTP3 > DTP1        202606  RESOLVED  2026-06-01  06-01\n'
    'DQ-MN33 Addi Arekay HC    Missing report     202605  RESOLVED  2026-06-01  06-01\n'
    '[> expand] Full SMS thread for DQ-4XDF:\n'
    '  2026-06-01 10:02  OUT  AHEAD DQ Alert [DQ-4XDF] Addi Arekay Health Center...\n'
    '  2026-06-01 10:04  IN   4\n'
    '  2026-06-01 10:04  OUT  [DQ-4XDF] Please reply with the correct value...\n'
    '  2026-06-01 10:05  IN   97\n'
    '  2026-06-01 10:05  OUT  [DQ-4XDF] Confirm change: BCG <1yr 970 -> 97...\n'
    '  2026-06-01 10:06  IN   YES\n'
    '  2026-06-01 10:06  OUT  [DQ-4XDF] Noted. Value corrected to 97 in DHIS2. Thank you.\n'
    '--------------------------------------------------------------------------------'
)
sp()

h2('Issue log columns')
tbl(
    ['Column', 'Content'],
    [
        ['Ref ID',          'e.g. DQ-4XDF — unique across all issues, never reused'],
        ['Facility',        'Facility name and woreda in parentheses'],
        ['Issue type',      'Outlier (element + value), DTP3 > DTP1, or Missing report'],
        ['Period',          'DHIS2 period code, displayed as human-readable month (Apr 2026)'],
        ['Status',          'OPEN | NOTIFIED | IN PROGRESS | CONFIRMING | RESOLVED | CONFIRMED OK | ESCALATED'],
        ['Cascade level',   'Current level: Facility / Woreda / Zone / Region / National'],
        ['Opened',          'Timestamp when issue was first detected'],
        ['Last contact',    'Timestamp of last outbound SMS'],
        ['Resolved value',  'What the value was corrected to (blank if not yet resolved)'],
        ['[Expand]',        'Click to see the full SMS thread for that issue'],
    ]
)
sp()

h2('Filters')
tbl(
    ['Filter', 'Options'],
    [
        ['Status',   'All / Open / Notified / In Progress / Confirming / Resolved / Confirmed OK / Escalated'],
        ['Type',     'All / Outlier / DTP inconsistency / Missing report'],
        ['Woreda',   'All / [woreda names from org unit hierarchy]'],
        ['Period',   'All / [YYYYMM dropdown]'],
        ['Search',   'Free-text search on ref ID or facility name'],
    ]
)
sp()

# ── 9. Edge Cases ─────────────────────────────────────────────────────────────
h1('9. Edge Cases and Agent Behaviour')
tbl(
    ['Edge case', 'Agent behaviour'],
    [
        ['Inbound SMS contains no reference ID',
         'Agent checks if the sender\'s phone number has exactly one open issue. If yes, routes to that issue. If multiple open issues, replies: "Please include your issue ref ID (e.g. DQ-4XDF) so we can route your reply."'],
        ['Reference ID not found or already closed',
         'Agent replies: "Ref ID not recognised or issue already closed. View all issues at <issue log URL>."'],
        ['Worker replies with non-numeric text when a number is expected',
         'Claude returns intent=unknown. Agent replies: "Please reply with a number (the option number or the correct value) as prompted."'],
        ['Worker replies NO at confirmation step',
         'Issue reverts to awaiting_option state. The original alert is re-sent verbatim so the worker can choose a different option.'],
        ['Worker sends a value that still looks like an outlier',
         'Agent applies the confirmed value as requested (after YES confirmation), then re-runs DQ checks in the next poll cycle. If the new value still fails, a new issue is created with a new ref ID.'],
        ['Same facility has two open issues simultaneously',
         'Both have distinct ref IDs. Every outbound message includes the ref ID. Worker must include the ref ID in replies — if missing, agent uses the single-open-issue routing logic above.'],
        ['Correction POST to DHIS2 fails (DHIS2 down or network error)',
         'Agent sends SMS: "Could not apply the change right now. Please re-enter the value directly in DHIS2 and reply DONE once submitted." Issue status remains CONFIRMING. Retry flagged for next cycle.'],
        ['Worker replies DONE (missing report already submitted)',
         'Agent schedules a re-check in 1 hour via APScheduler. If the registration appears, issue closes automatically with a confirmation SMS. If not, issue remains open.'],
        ['Escalation notification itself fails to deliver',
         'Twilio delivery failure is logged. Agent retries the escalation SMS once after 30 minutes. Failure is visible on the issue log page under last_contact.'],
        ['Issue resolved at woreda level (not facility)',
         'If woreda officer contacts AHEAD team directly and they manually close the issue via the log page, status is set to resolved with a note. No further SMS is sent.'],
    ]
)
sp()

# ── 10. Reference ID Design ───────────────────────────────────────────────────
h1('10. Reference ID Design')
body('Every issue gets a unique ID on creation. The ID appears in every outbound message and is how the agent routes inbound replies to the right issue. Design constraints:')
sp()
bullet('Visually unambiguous when read aloud or off a small phone screen')
bullet('Short enough to quote in an SMS without taking up too much space')
bullet('Random enough to be hard to guess (prevents spoofed replies)')
bullet('Never reused — resolved issues keep their ID permanently for audit')
sp()
tbl(
    ['Property', 'Decision', 'Reason'],
    [
        ['Format', 'DQ-XXXX (4-character suffix)', 'DQ prefix identifies the source system; 4 chars is short and unambiguous'],
        ['Character set',
         'A B C D E F G H J K M N P Q R S T U V W X Y + 3 4 5 6 7 8 9 (29 chars)',
         'Excludes 0/O, 1/I/L, 2/Z — characters that look alike on paper or screen'],
        ['Capacity', '29^4 = 707,281 unique IDs', 'More than sufficient for MVP; Phase 2 can extend to 5 chars if needed'],
        ['Generation', 'Python secrets.choice() from the charset', 'Cryptographically random — hard to guess or enumerate'],
        ['Collision check', 'DB lookup before issuing', 'Probability of collision is < 0.001% even at 10,000 open issues'],
    ]
)
sp()

# ── 11. WhatsApp Mode ──────────────────────────────────────────────────────────
h1('11. WhatsApp Mode (Demo / Production Toggle)')
body('For the demo, the agent uses the Twilio WhatsApp sandbox. Set TWILIO_WHATSAPP=true in .env to activate WhatsApp mode. Set TWILIO_WHATSAPP=false for standard SMS in production.')
sp()
tbl(
    ['Mode', 'TWILIO_WHATSAPP value', 'How it works'],
    [
        ['WhatsApp sandbox (demo)',
         'true',
         'Messages sent and received via the Twilio WhatsApp sandbox number (+1 415 523 8886). The presenter\'s phone must first send "join [keyword]" to that number to opt in to the sandbox.'],
        ['Standard SMS (production)',
         'false',
         'Messages sent and received via a standard Twilio SMS number. No opt-in required. Works on any phone including basic feature phones.'],
    ]
)
sp()

# ── 12. MVP vs Phase 2 Scope ──────────────────────────────────────────────────
h1('12. MVP vs Phase 2 Scope')
body('The following table clarifies what is in scope for this build and what is explicitly deferred. Phase 2 items are excluded not because they are unimportant but because they either require external dependencies (Meta WhatsApp approval, DHIS2 App Framework expertise) or can be validated after the core SMS loop is working.')
sp()
tbl(
    ['Feature', 'MVP', 'Phase 2', 'Notes'],
    [
        ['Post-commit DQ checks (polling)',                     'Yes', '—',  ''],
        ['Outlier detection: Z-score > 2.0 SD (AHEAD Method 1)', 'Yes', '—', 'AHEAD Table 1 operational threshold'],
        ['Outlier detection: absolute diff > 100 doses (AHEAD Method 5)', 'Yes', '—', ''],
        ['DTP1/DTP3 consistency check (30% relative or 100 absolute)', 'Yes', '—', 'Monthly facility threshold per AHEAD guide'],
        ['Missing report detection',                            'Yes', '—',  ''],
        ['Two-way SMS with YES/NO confirmation before write-back', 'Yes', '—', 'All DHIS2 write-back options require confirmation'],
        ['Full escalation cascade (facility → national)',       'Yes', '—',  ''],
        ['Apply corrections back to DHIS2 via API',            'Yes', '—',  ''],
        ['Issue log web page (standalone Flask, auto-refresh)', 'Yes', '—',  'Dashboard at localhost:5001/issues'],
        ['Claude API (reply parsing + message generation)',     'Yes', '—',  ''],
        ['English-language messages only',                      'Yes', '—',  ''],
        ['5-method outlier ensemble (R script integration)',    '—',  'Yes', 'Requires access to AHEAD R scripts; deferred'],
        ['Name consistency check',                              '—',  'Yes', 'Not covered by DHIS2 built-in rules'],
        ['WhatsApp channel (production / approved)',            '—',  'Yes', 'Requires Meta Business API approval (1-4 weeks); sandbox available for demo'],
        ['Email channel',                                       '—',  'Yes', 'Reply parsing noisier; lower priority'],
        ['Pre-commit inline form warning (custom form JS)',     '—',  'Yes', 'Requires DHIS2 custom form; useful but not needed to validate cascade'],
        ['Amharic messages',                                    '—',  'Yes', 'Claude can generate; need translation review before sending'],
        ['DHIS2 embedded app (issue log inside DHIS2)',        '—',  'Yes', 'Requires DHIS2 App Framework (React); standalone page sufficient for MVP'],
        ['Inbound phone call routing',                          '—',  'Yes', 'Higher complexity; SMS covers facility workers adequately for MVP'],
    ]
)
sp()

# ── 13. Demo Script ───────────────────────────────────────────────────────────
h1('13. Demo Script')
body('Three live data-entry demos, all using the same facility and user account: eth_facility_01 (Almaz Tadesse), Addi Arekay Health Center. You type wrong values directly into the DHIS2 form as the facility worker would, the agent detects the issue within ~30 seconds, and the numbered-option SMS conversation plays out on the presenter\'s phone. Total runtime: ~15 minutes.')
sp()

tbl(
    ['Demo', 'Issue type', 'Facility', 'User', 'Period', 'Duration', 'What it proves'],
    [
        ['A', 'Statistical outlier',     'Addi Arekay HC', 'eth_facility_01', 'Jun 2026', '~6 min', 'Type BCG=970 (meant 97) → outlier detected → option 4 → value entered → confirmed YES → corrected to 97 in DHIS2'],
        ['B', 'DTP1/DTP3 inconsistency', 'Addi Arekay HC', 'eth_facility_01', 'Jun 2026', '~5 min', 'Edit Penta3=90 > Penta1=60 → validation rule → option 2 → confirmed YES → both corrected'],
        ['C', 'Missing report',          'Addi Arekay HC', 'eth_facility_01', 'May 2026', '~4 min', 'No data for prior month → agent flags → SUBMIT → recovery flow → auto-close'],
        ['D', 'Issue log web app',       '(all)',           '—',               '(all)',    '~3 min', 'Supervisor view: all issues, full 5-turn threads with confirmation step, statuses'],
    ]
)
sp()

h2('Prerequisites — confirm before presenting')
bullet('DHIS2 running: http://localhost:8080 responds with login page')
bullet('Agent running: http://localhost:5001/api/status returns JSON with uptime')
bullet('inject_data.py has been run — 28 months of clean baseline data loaded (Jan 2024 – Apr 2026). No pre-seeded issues. Addi Arekay Health Center has 22 months of clean baseline values.')
bullet('contacts table in agent SQLite seeded with presenter\'s phone for Addi Arekay Health Center')
bullet('Presenter\'s phone visible to audience (screen mirror or camera pointing at phone)')
bullet('Twilio account active and TWILIO_FROM_NUMBER configured in .env')
bullet('TWILIO_WHATSAPP=true in .env (sandbox mode for demo) — presenter\'s phone has sent "join [keyword]" to +1 415 523 8886')
sp()

body('All demos use a single login: eth_facility_01 / [DHIS2_USER_PASSWORD from .env]. This account belongs to Almaz Tadesse and is scoped to Addi Arekay Health Center only — exactly the facility-worker persona being demonstrated.')
sp()
body('Demos A and B use June 2026 (202606) — a period with no data yet. Demo C uses May 2026 (202605) to show the missing report scenario for a prior month that was never formally completed.')
sp()
body('Addi Arekay Health Center is a Health Center (tier 2). BCG under-1 baseline: ~97/month (range 80–115 across 22 months). Penta1 under-1: ~87. Penta3 under-1: ~71. Outlier detection threshold: Z-score > 2.0 SD OR absolute diff > 100 doses. Detection is within ~30 seconds via lastUpdated polling, or trigger instantly with: POST http://localhost:5001/api/scan')
sp()

h2('Demo A — Statistical Outlier: Addi Arekay Health Center, Jun 2026 (~6 min)')
body('What this shows: Almaz enters her June 2026 BCG count with an extra zero — a realistic typo. The agent detects it against 22 months of baseline history (Z-score > 2.0 SD threshold) and sends the numbered-option SMS. The full loop — including the YES/NO confirmation step — plays out in front of the audience in real time.')
sp()

body('Step 1 — Log in and open the data entry form')
bullet('Log in at http://localhost:8080 as eth_facility_01 / [DHIS2_USER_PASSWORD from .env]')
bullet('Open Data Entry (grid icon in top menu)')
bullet('Navigate in the left org unit tree: Ethiopia → Amhara → North Gondar → Addi Arekay Woreda → Addi Arekay Health Center')
bullet('Select dataset: EPI - Routine vaccine delivery | Period: June 2026')
bullet('The form is blank — no data for June 2026 yet')
bullet('Narrate: "Almaz is entering the monthly EPI report for June. She works at the facility and submits this form each month."')
sp()

body('Step 2 — Enter data with a BCG outlier (deliberate typo)')
bullet('Enter the following values in the form:')
bullet('BCG under-1: 970  |  BCG >= 1 year: 5', lv=1)
bullet('Penta1 under-1: 88  |  Penta1 >= 1 year: 4', lv=1)
bullet('Penta3 under-1: 72  |  Penta3 >= 1 year: 3', lv=1)
bullet('MR1 under-1: 83  |  MR1 >= 1 year: 3', lv=1)
bullet('Narrate: "BCG under-1 should be 97 this month — but Almaz accidentally typed 970. An extra zero. This happens."')
bullet('Click Save')
sp()

body('Step 3 — Agent detects the outlier (~30 seconds)')
bullet('Wait ~30 seconds, OR trigger immediately: curl -s -X POST http://localhost:5001/api/scan')
bullet('Agent runs outlier detection: BCG = 970 vs 22-month mean ~97. Z-score well above 2.0 SD threshold — flagged. Also exceeds 100-dose absolute diff threshold.')
bullet('New issue created. Open http://localhost:5001/issues — issue appears as NOTIFIED')
sp()

body('Step 4 — SMS arrives on presenter\'s phone')
mono(
    'AHEAD DQ Alert [DQ-XXXX]\n'
    'Addi Arekay Health Center — Jun 2026\n'
    'BCG: 970 doses (expected 79–114)\n'
    'Reply with option number:\n'
    '1. Replace with 6-month average\n'
    '2. Keep as-is (no action)\n'
    '3. Set to zero\n'
    '4. Replace with specific value\n'
    '5. At health facility doses only\n'
    '6. Outreach doses only'
)
bullet('Point out: numbered options mirror the AHEAD reference guide (section 2.2) exactly — same decision schema George\'s team uses today, now delivered by SMS')
bullet('Issue log status is now IN PROGRESS once the first reply arrives')
sp()

body('Step 5 — Select option 4 (replace with specific value)')
bullet('Reply to the SMS: 4')
bullet('Agent sends a follow-up:')
mono(
    '[DQ-XXXX] Please reply with the correct value\n'
    '(numbers only, e.g. 97).'
)
bullet('Reply: 97')
bullet('Agent sends the YES/NO confirmation (issue log status: CONFIRMING):')
mono(
    '[DQ-XXXX] Confirm change:\n'
    'BCG <1yr — Addi Arekay Health Center (Jun 2026)\n'
    '970 -> 97\n'
    'Reply YES to update DHIS2 or NO to re-enter.'
)
bullet('Reply: YES')
bullet('Agent applies the correction and sends:')
mono(
    '[DQ-XXXX] Noted. Value corrected to 97 in DHIS2. Thank you.'
)
bullet('Issue log status changes to RESOLVED')
sp()

body('Step 6 — Show the correction in DHIS2')
bullet('Switch back to DHIS2 and refresh Addi Arekay HC / June 2026')
bullet('BCG under-1 now shows 97 (was 970)')
bullet('Narrate: "The correction was written to DHIS2 by the agent after explicit YES confirmation. The audit log will show agent_service made this change — fully traceable."')
sp()

body('Expected final state:')
bullet('DHIS2: Addi Arekay HC, Jun 2026, BCG under-1 = 97 (was 970)')
bullet('Issue log: DQ-XXXX | Outlier: BCG | Addi Arekay HC | RESOLVED | Option 4: specific value (97)')
bullet('SMS thread: 5 messages — notification, "4", follow-up question, "97", confirmation request, "YES", resolution (5 turns visible in log)')
sp()

h2('Demo B — DTP1/DTP3 Inconsistency: Addi Arekay Health Center, Jun 2026 (~5 min)')
body('What this shows: Almaz accidentally enters a higher value for Penta3 (DTP3) than for Penta1 (DTP1). This is biologically impossible — more children cannot have completed the vaccine series than started it. The agent detects the gap (50% relative difference, above the 30% AHEAD threshold), sends an alert with the same response options as the AHEAD Excel dropdown, and automatically corrects the value in DHIS2 after confirmation.')
sp()

body('Step 1 — Navigate to the same June 2026 form')
bullet('Still logged in as eth_facility_01.')
bullet('Data Entry → Addi Arekay Health Center → EPI - Routine vaccine delivery → June 2026')
sp()

body('Step 2 — Enter the inconsistent values')
bullet('In the Vaccinations - children table:')
bullet('DPT-HepB-HIB 1, < 1 year: type 60, press Tab')
bullet('DPT-HepB-HIB 3, < 1 year: type 90, press Tab')
bullet('Both fields turn green (saved).')
bullet('Narrate: "Almaz copied the figures from her paper tally sheet but swapped the rows. DTP3 = 90 > DTP1 = 60 is epidemiologically impossible — dropout only flows one direction."')
bullet('Threshold: |(90-60)/60| = 50% relative gap, above the 30% AHEAD facility-monthly threshold.')
sp()

body('Step 3 — Watch the issue log (~30 seconds)')
bullet('At http://localhost:5001/issues a new row appears: Type=DTP, Facility=Addi Arekay HC, Value=90, Status=NOTIFIED')
sp()

body('Step 4 — WhatsApp alert arrives')
mono(
    'AHEAD DQ Alert [DQ-YYYY]\n'
    'Addi Arekay Health Center — Jun 2026\n'
    'DTP1=60, DTP3=90 (DTP3 > DTP1, gap: 50%)\n'
    '\n'
    'Reply with option number:\n'
    '1. Keep as-is (no action)\n'
    '2. Use DTP1 value for both\n'
    '3. Use DTP3 value for both\n'
    '4. Replace with specific value\n'
    '5. Other reason'
)
bullet('Narrate: "The options match the AHEAD Excel dropdown exactly — same schema, just delivered by WhatsApp instead of email."')
sp()

body('Step 5 — Select option 2 (use DTP1 for both)')
bullet('Narrate: "Almaz knows Penta1 = 60 is correct. She selects option 2 — agent will set DTP3 to match DTP1."')
bullet('Reply: 2')
bullet('Issue log updates to CONFIRMING. Agent sends:')
mono(
    '[DQ-YYYY] Confirm: Set DTP3 to match DTP1 (60)\n'
    'Addi Arekay Health Center (Jun 2026)\n'
    '\n'
    'Reply YES to update DHIS2 or NO to choose again.'
)
sp()

body('Step 6 — Confirm')
bullet('Reply: YES')
bullet('Agent writes Penta3 = 60 to DHIS2. Issue log updates to RESOLVED. Confirmation:')
mono(
    '[DQ-YYYY] Noted. DTP3 set to match DTP1 (60) in DHIS2. Thank you.'
)
sp()

body('Step 7 — Verify in DHIS2')
bullet('Refresh the Addi Arekay HC / June 2026 form.')
bullet('DPT-HepB-HIB 3, < 1 year now shows 60 (updated by agent_service, matching DTP1).')
sp()

body('Expected final state:')
bullet('DHIS2: DTP1=60, DTP3=60 — Penta3 corrected to match Penta1')
bullet('Issue log: RESOLVED — "DTP3 set to match DTP1 (60) in DHIS2."')
bullet('SMS thread: 3 turns — alert / "2" / confirmation / "YES" / resolution')
sp()

h2('Demo C — Missing Report: Addi Arekay Health Center, May 2026 (~4 min)')
body('What this shows: The agent proactively flags a prior month (May 2026) where no complete dataset registration exists for Addi Arekay HC. This is a different scenario from Demos A and B — rather than catching errors at entry time, it catches a facility that never submitted at all. The AHEAD methodology is explicit: recovery first, imputation only as a fallback.')
sp()

body('Step 1 — Show that May 2026 has no submission')
bullet('Navigate to Addi Arekay HC | Dataset: EPI - Routine vaccine delivery | Period: May 2026')
bullet('The form has data values (loaded by inject_data.py) but no complete registration — the "Complete" button was never clicked')
bullet('Narrate: "In the real workflow, the agent checks which facilities have submitted and formally completed their report. Addi Arekay HC has data values for May 2026 but has never been formally completed — the HMIS officer never clicked the Complete button. The agent flags this."')
sp()

body('Step 2 — Trigger the missing report check')
bullet('POST http://localhost:5001/api/scan?check=missing_reports')
bullet('Agent queries completeDataSetRegistrations for May 2026; Addi Arekay HC has no registration')
bullet('Issue created: type=missing_report, facility=Addi Arekay HC, period=May 2026')
sp()

body('Step 3 — SMS arrives on presenter\'s phone')
mono(
    'AHEAD DQ Alert [DQ-XXXX]\n'
    'Addi Arekay Health Center\n'
    'No completed EPI report for May 2026.\n'
    'Can you submit the report? Reply SUBMIT.\n'
    'Or if data is unavailable, select:\n'
    '1. Will submit by [date]\n'
    '2. Data cannot be recovered\n'
    '3. Facility closed / no service that month\n'
    '4. Other'
)
bullet('Point out: SUBMIT appears first — the AHEAD methodology says recovery beats imputation; only choose a cleaning option if the report genuinely cannot be found. Options 2 and 3 are flagged for HQ zero imputation rather than being auto-written.')
sp()

body('Step 4 — Reply SUBMIT')
bullet('Reply: SUBMIT')
bullet('Agent sends:')
mono(
    '[DQ-XXXX] Thanks — checking DHIS2 within the hour.\n'
    'Issue will close automatically once your\n'
    'submission appears.'
)
bullet('Now in DHIS2: navigate to Addi Arekay HC / May 2026, click the Complete button to mark it as formally submitted')
bullet('Trigger a re-check: POST http://localhost:5001/api/scan?check=missing_reports')
bullet('Agent finds the registration and sends:')
mono(
    '[DQ-XXXX] Confirmed. Addi Arekay HC May 2026\n'
    'report received. Issue closed.'
)
sp()

body('Expected final state:')
bullet('DHIS2: Addi Arekay HC, May 2026 — marked as complete (Complete button clicked)')
bullet('Issue log: DQ-XXXX | Missing report | Addi Arekay HC | RESOLVED')
bullet('SMS thread: 3 messages — notification, "SUBMIT", acknowledgement, auto-close confirmation')
sp()

h2('Demo D — Issue Log Web App (~3 min)')
body('What this shows: The standalone supervisor view. George\'s team uses this to monitor all DQ activity across Ethiopia — every issue the agent detected, every SMS sent, every decision made — without ever logging into DHIS2. The dashboard auto-refreshes every 10 seconds.')
sp()

body('Step 1 — Open the issue log')
bullet('Navigate to http://localhost:5001/issues in the browser')
bullet('Three issues visible: Demo A (RESOLVED, BCG outlier), Demo B (RESOLVED, DTP inconsistency), Demo C (RESOLVED, missing report)')
sp()

body('Step 2 — Show filtering')
bullet('Filter by Status = "Open" — shows only unresolved issues (none after all three demos resolved)')
bullet('Filter by Facility = "Addi Arekay HC" — shows all three issues for this facility')
bullet('Narrate: "In production with hundreds of facilities across Ethiopia, George\'s team could filter by any woreda, region, or time period."')
sp()

body('Step 3 — Expand Demo A conversation thread (BCG outlier)')
bullet('Click expand on the BCG outlier issue (DQ-XXXX)')
bullet('Show the full SMS thread: timestamps, IN/OUT, each turn')
bullet('The thread shows: initial notification → "4" (option selected) → follow-up question → "97" (value given) → confirmation request → "YES" → resolution')
bullet('Narrate: "This is the audit trail. It captures the same information the Excel dropdown captured — which option was chosen, what the corrected value is — but in real time, with the full conversation history including the explicit confirmation step."')
sp()

body('Step 4 — Point out the status progression column')
bullet('Issue record shows status: RESOLVED, and the status history: OPEN → NOTIFIED → IN PROGRESS → CONFIRMING → RESOLVED')
bullet('Narrate: "The CONFIRMING state is new — it represents the moment between the worker providing a value and actually writing to DHIS2. No data is changed until YES is received. That\'s the key safety guarantee."')
sp()

body('Expected final state:')
bullet('Audience has seen the complete loop: facility entry → real-time detection → numbered SMS options → value entry → YES/NO confirmation → correction applied → audit log with full thread')
bullet('Issue log shows 3 resolved issues for Addi Arekay HC across two periods, with decision records and status progression')
sp()

h2('Common demo questions and answers')
tbl(
    ['Question', 'Answer'],
    [
        ['Why is there a YES/NO confirmation step?',
         'No data is written to DHIS2 until the worker explicitly confirms. This prevents the agent from applying a value the worker typed by accident and gives a clear moment in the audit trail where the human approved the change.'],
        ['Why numbered options instead of free text?',
         'These are the same fixed categories in the current AHEAD reference guide (sections 2.2–2.4). Using them means the agent captures exactly the same decisions the manual process captures, in a format the AHEAD team already understands.'],
        ['What if the facility worker ignores the SMS?',
         'The agent retries every 24h for 3 days. After 72h with no response, it automatically sends a summary to the woreda HMIS officer. No manual escalation needed.'],
        ['Who maintains the phone number list?',
         'The AHEAD team maintains a contacts registry (a simple table mapping each org unit to a phone number). It is loaded into the agent at startup and updated as staff change.'],
        ['Does the facility worker need to install an app?',
         'No. Standard SMS only. Works on any phone, including basic feature phones without internet. For the demo, the Twilio WhatsApp sandbox is used; production uses standard SMS.'],
        ['Can the agent send in Amharic?',
         'Not in MVP. Claude can generate Amharic and this is a Phase 2 upgrade pending translation review.'],
        ['What if the agent applies the wrong value?',
         'The agent logs every option selection, every confirmed value, and the exact write-back applied. The confirmation step means the facility worker saw and approved the exact change. The issue thread is a complete audit trail. If a mistake is made, the AHEAD team can see it in the log and manually correct in DHIS2.'],
        ['What happens to issues that never get resolved?',
         'They escalate automatically through the hierarchy (facility → woreda → zone → region → national digest) and remain open in the log indefinitely until closed.'],
        ['What DQ methods does the MVP implement?',
         'Two of the five AHEAD methods: SD-based outlier detection (Method 1, Z-score > 2.0 SD) and absolute difference checks (Method 5, > 100 doses). DTP consistency uses a 30% relative or 100 absolute dose threshold per the AHEAD guide.'],
    ]
)
sp()

out = HERE / 'AHEAD_AI_UX_Workflow.docx'
doc.save(str(out))
print(f'Saved: {out}')
