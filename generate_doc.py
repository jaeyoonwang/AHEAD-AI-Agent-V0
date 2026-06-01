#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
    p.runs[0].font.size = Pt(15)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
    p.runs[0].font.size = Pt(12)
    return p

def body(text):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.size = Pt(10.5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def mono(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'D6E4F0')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            if cells[c_idx].paragraphs[0].runs:
                cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
    return t

def spacer():
    doc.add_paragraph('')

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('AHEAD AI Agent — MVP Architecture', 0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
title.runs[0].font.size = Pt(18)

sub = doc.add_paragraph('UNICEF AHEAD × Gates Foundation AI Fellows  |  May 2026')
sub.runs[0].font.size = Pt(9)
sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
sub.runs[0].italic = True
spacer()

# ── 1. Problem ────────────────────────────────────────────────────────────────
h1('1. The Problem')
body("The current DQ pipeline is fully manual: R scripts run offline on UNICEF laptops, output flagged Excel files, which then kick off an email and phone cascade up and down the hierarchy until corrections are made. Not only does this make it extremely difficult to pinpoint where the process is at a given moment or to track and follow up in a timely manner, it also doesn't scale and leaves the door open for errors to go unresolved for weeks. The goal of the AI agent is to automate this loop — catch errors close to the source, notify the right person immediately, and only escalate what genuinely can't be resolved at the facility level.")
spacer()

h2('Current workflow (as-is)')
body('The diagram below shows the existing manual process. The AI agent replaces steps 2–12 — the Excel production, email cascade, and manual phone outreach — with an automated, real-time notification and correction loop.')
spacer()

from docx.shared import Inches as _Inches
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture('ahead_pipeline_diagram.png', width=_Inches(5.5))
spacer()

# ── 2. Workflow ───────────────────────────────────────────────────────────────
h1('2. Where the AI Agent Fits')
body('The agent runs post-commit — after data lands in DHIS2 — and works in two parallel modes: it checks new submissions as they come in, and separately runs a daily scan for facilities that never submitted at all.')
spacer()

mono(
    'Facility submits EPI form\n'
    '         │\n'
    '         ▼\n'
    'DHIS2 datavalue table   ← data written immediately\n'
    '         │\n'
    '         ▼  API poll every 5 min\n'
    ' ┌────────────────────┐      ┌─────────────────────────────┐\n'
    ' │   DQ ENGINE        │      │  DAILY CRON                 │\n'
    ' │  (on new submit)   │      │  (missing report check)     │\n'
    ' └────────┬───────────┘      └──────────────┬──────────────┘\n'
    '          │ issues found?                    │ facility did not submit?\n'
    '          └──────────────┬───────────────────┘\n'
    '                         ▼\n'
    '             Assign reference ID (e.g. DQ-AX42)\n'
    '             Notify facility via SMS\n'
    '                         │\n'
    '                         ▼\n'
    '          ┌──────────────────────────────┐\n'
    '          │   CASCADE STATE MACHINE      │\n'
    '          │   tracks each open issue     │\n'
    '          └──────────────┬───────────────┘\n'
    '                         │\n'
    '           resolved? ────┤──── 3 days unresolved?\n'
    '               ↓         │              ↓\n'
    '             close        │      escalate → woreda\n'
    '                          │              │\n'
    '                          │      1 week unresolved?\n'
    '                          │              ↓\n'
    '                          │      escalate → region\n'
    '                          │              │\n'
    '                          │      end of month?\n'
    '                          │              ↓\n'
    '                          │      national summary\n'
    '                          │\n'
    '         all events logged to Data Issue Log (standalone web page)'
)
spacer()

body('One important nuance on the cascade: catching errors at submission time will reduce escalation volume significantly, but it does not eliminate the cascade entirely. Three cases still require escalation: (1) missing reports — if a facility never submits, there is nothing to check; (2) dismissed warnings — a facility worker can acknowledge the flag and submit anyway; (3) systemic patterns — some anomalies are only visible when looking across multiple facilities at the woreda or zone level, not from a single submission.')
spacer()

# ── 3. DQ Engine ──────────────────────────────────────────────────────────────
h1('3. DQ Engine')

h2('Origin of the four checks')
body("These four checks are not from the EPI metadata package — they come from AHEAD's existing DQ methodology, currently implemented as R scripts that run offline. The EPI package does have its own built-in validation rules (the 'Run Validation' button in DHIS2 Data Entry), but those are more basic plausibility checks within a single form. What we're automating here is the broader, cross-facility analysis the AHEAD team already runs manually.")
spacer()

add_table(
    ['Check', 'DHIS2 API endpoint', 'Flag condition'],
    [
        ['Missing report',
         'GET /api/completeDataSetRegistrations',
         'Facility has no completed submission by day 5 of the following month'],
        ['Statistical outlier',
         'GET /api/outlierDetection (Z-score + IQR)',
         'DHIS2 built-in outlier detection flags value as anomalous vs. facility historical baseline. Threshold configurable (default z=3.0). Full 5-method ensemble (SD, MAD, Lowess, etc.) deferred to Phase 2.'],
        ['DTP1/DTP3 consistency',
         'GET /api/validationResults',
         'EPI metadata package includes a built-in validation rule: Penta3 ≤ Penta1. Any violation is returned directly by this endpoint.'],
        ['Name consistency',
         'GET /api/organisationUnits + dataValueSets',
         'Compare submitted org unit codes against the reference hierarchy. Deferred to Phase 2 — not covered by DHIS2 built-in rules.'],
    ]
)
spacer()

h2('Why DHIS2 API-native checks for MVP')
body('Rather than calling external R scripts via subprocess, the MVP uses DHIS2\'s own REST API for all DQ checks. This is the right approach for Phase 1 because:')
bullet('No external dependencies — no R installation, no script access required; agent is a single Python process')
bullet('The EPI Aggregate Metadata Package already ships with DTP1/DTP3 validation rules; no reimplementation needed')
bullet('DHIS2\'s outlier detection endpoint is sufficient to catch 10× spikes; calibrated thresholds are a Phase 2 concern')
bullet('Reduces the critical path — no dependency on obtaining and reviewing the AHEAD team\'s R scripts before building can start')
spacer()

mono(
    'DHIS2 API  →  GET /api/completeDataSetRegistrations  →  missing report flags\n'
    '           →  GET /api/outlierDetection              →  outlier flags\n'
    '           →  GET /api/validationResults             →  DTP consistency flags\n'
    '                         │\n'
    '               Python consolidates results\n'
    '                         │\n'
    '               Creates issue objects  →  cascade'
)
spacer()
body('The full AHEAD 5-method ensemble (SD, MAD, Median AD, Lowess, Absolute diff from mean) is the Phase 2 upgrade path, once access to the existing R scripts is available. Swapping in the R subprocess wrapper at that point requires only changing the DQ engine module — nothing else in the stack changes.')
spacer()

# ── 4. Two-way SMS conversation ───────────────────────────────────────────────
h1('4. Two-Way SMS Conversation Loop')

h2('Communication channel decision')
body('For MVP, we are starting with SMS only. WhatsApp is the more natural channel for most users but requires Meta Business API approval which can take 1–4 weeks and is a real timeline risk. Email reply parsing is also noisier in practice (quoted text, auto-signatures, threading). SMS works immediately with no approval process, works on basic phones without internet — which matters for rural facility workers — and the forced brevity is actually useful here. The conversation logic itself is channel-agnostic, so WhatsApp and email can be added in Phase 2 without changing any of the underlying state machine.')
spacer()

h2('Reference IDs')
body('Every issue gets a short reference ID when it is created (e.g. DQ-AX42). The ID only uses characters that are visually unambiguous — no 0/O, 1/I/L, 2/Z — so there is no confusion when reading it off a phone screen. It appears in every outbound message and is used to match inbound replies back to the right issue. This solves two problems cleanly:')
bullet('If a facility has multiple open issues, the reference ID makes clear which one is being addressed')
bullet('It acts as a lightweight token — random enough to be hard to guess, which prevents spoofed replies')
spacer()

h2('Conversational confirmation flow')
body('The agent does not apply corrections automatically. It proposes the change and asks for confirmation first. If the facility worker gives a different number than the agent expected, or corrects the agent mid-conversation, the agent adapts without requiring them to start over. The goal is for this to feel like a natural back-and-forth, not a rigid command interface.')
spacer()

mono(
    'Agent  →  "[ DQ-AX42 ] Limalimo HP, Apr 2026: BCG <1yr = 350\n'
    '           (normal range 30-45, ~10x high). Was this an error?\n'
    '           Reply with the correct value or KEEP to leave as-is."\n'
    '\n'
    'Facility →  "35"\n'
    '\n'
    'Agent  →  "[ DQ-AX42 ] Confirm: change BCG <1yr at Limalimo HP\n'
    '           (Apr 2026) from 350 to 35.\n'
    '           Reply YES to apply or NO to cancel."\n'
    '\n'
    'Facility →  "no wait, should be 53"\n'
    '\n'
    'Agent  →  "[ DQ-AX42 ] Got it — so the correct value is 53?\n'
    '           Reply YES to apply 53 or send the correct number."\n'
    '\n'
    'Facility →  "YES"\n'
    '\n'
    'Agent  →  "[ DQ-AX42 ] Done — BCG <1yr updated to 53.\n'
    '           Issue resolved. ✓"'
)
spacer()

body('An LLM parses inbound replies to extract the intended corrected value from free text. This handles messages like "oh that was a typo, should be 53" without requiring the facility worker to follow a structured format. Hallucination risk is low here because the task is extracting a specific number from a very bounded context — it is not open-ended generation.')
spacer()

add_table(
    ['Conversation state', 'Trigger', 'Agent action'],
    [
        ['notified',         'Issue created',                           'Send initial SMS with reference ID'],
        ['awaiting_confirm', 'Facility replies with a value',           'Propose specific change, ask YES/NO'],
        ['resolved',         'Facility replies YES',                    'Apply fix to DHIS2 via API, close issue, log'],
        ['dismissed',        'Facility replies KEEP',                   'Leave value as-is, close issue, log'],
        ['retry',            'No reply after 24h (×3)',                 'Re-send notification'],
        ['escalated',        '72h no resolution',                       'Hand off to woreda HMIS officer'],
    ]
)
spacer()

# ── 5. Data Issue Log ─────────────────────────────────────────────────────────
h1('5. Data Issue Log (Standalone Web Page)')
body('All issues, SMS conversations, and resolution actions are logged and viewable on a standalone web page. It is separate from DHIS2 — it reads from the agent\'s own database — but would be the primary place for supervisors and the AHEAD team to monitor what the agent is doing and intervene if needed.')
spacer()

add_table(
    ['Column', 'Description'],
    [
        ['Ref ID',             'e.g. DQ-AX42'],
        ['Issue type',         'outlier / missing report / DTP inconsistency'],
        ['Facility',           'Facility name and woreda'],
        ['Period',             'e.g. Apr 2026'],
        ['Flagged value',      'The value that triggered the check'],
        ['Resolved value',     'What it was corrected to (if resolved)'],
        ['Status',             'notified / awaiting_confirm / resolved / escalated'],
        ['Conversation',       'Full SMS thread (expandable)'],
        ['Opened / Resolved',  'Timestamps'],
    ]
)
spacer()
body('Embedding this as a proper DHIS2 app (React, DHIS2 App Framework) is a Phase 2 improvement. A standalone page is fine for MVP — the data is what matters, not where it lives.')
spacer()

# ── 6. Technical stack ────────────────────────────────────────────────────────
h1('6. Technical Stack')

add_table(
    ['Layer', 'Technology', 'Rationale'],
    [
        ['DQ checks',          'DHIS2 REST API (outlierDetection, validationResults, completeDataSetRegistrations)',
         'MVP uses built-in endpoints — no external dependencies. Phase 2 upgrades to AHEAD R-script ensemble via subprocess wrapper.'],
        ['Agent orchestration', 'Python',
         'Glue layer: poll DHIS2, consolidate DQ results, manage state, send SMS, apply fixes'],
        ['LLM',                'Claude API (claude-haiku-4-5 for parsing, sonnet for generation)',
         'Parse free-text replies; generate context-aware notification messages'],
        ['DHIS2 integration',  'DHIS2 REST API only (GET dataValueSets, POST dataValueSets)',
         'No DHIS2 internals modified; agent reads and writes via the same API as any user'],
        ['State + audit log',  'SQLite (MVP) → PostgreSQL (production)',
         'Tracks issue lifecycle and full conversation history; survives restarts'],
        ['SMS',                'Twilio',
         'SMS only for MVP; WhatsApp and email are Phase 2 (same logic, different transport)'],
        ['Inbound webhook',    'Flask (receives Twilio SMS webhook)',
         'Routes inbound SMS to the correct open issue via reference ID'],
        ['Issue log page',     'Flask (standalone web page)',
         'Displays all issues, statuses, and conversation threads'],
        ['Scheduler',          'APScheduler',
         'Daily missing report cron; retry timers; monthly summary'],
        ['Deployment',         'Docker Compose (same host as DHIS2)',
         'No new infrastructure; agent container added to existing compose file'],
    ]
)
spacer()

# ── 7. Feasibility ────────────────────────────────────────────────────────────
h1('7. Feasibility Assessment')

add_table(
    ['Component', 'Effort', 'Risk', 'Notes'],
    [
        ['DQ engine (DHIS2 built-in checks)', 'Low',        'Low',    'API calls only — outlierDetection, validationResults, completeDataSetRegistrations. No R dependency.'],
        ['Missing report detection',          'Low',        'Low',    'completeDataSetRegistrations endpoint; single query per period'],
        ['Cascade state machine',             'Medium',     'Low',    'Standard pattern; main work is timer and retry logic'],
        ['Outbound SMS (Twilio)',              'Low',        'Low',    'Twilio SMS is straightforward; live in minutes'],
        ['Inbound SMS webhook + ref ID match','Low–Medium', 'Low',    'Flask webhook + reference ID routing is simple and reliable'],
        ['LLM reply parsing',                 'Low',        'Low',    'Bounded extraction task; low hallucination risk'],
        ['Apply fix to DHIS2',                'Low',        'Low',    'POST /api/dataValueSets with corrected value'],
        ['Standalone issue log page',         'Low',        'Low',    'Simple Flask page over SQLite; no DHIS2 app framework needed'],
        ['WhatsApp / email (Phase 2)',         'Medium',     'Medium', 'Meta approval takes weeks; email reply parsing is noisier'],
        ['Pre-commit DHIS2 app (Phase 2)',     'High',       'Medium', 'Requires custom React app in DHIS2 App Framework'],
    ]
)
spacer()

# ── 8. MVP scope ──────────────────────────────────────────────────────────────
h1('8. MVP Scope')

h2('In scope')
bullet('Post-commit DQ checks via DHIS2 REST API: outlier detection (outlierDetection endpoint), DTP1/DTP3 consistency (validationResults), missing report (completeDataSetRegistrations). Name consistency deferred to Phase 2.')
bullet('Full cascade: facility → woreda → zone → region → national')
bullet('Two-way SMS with reference IDs and a conversational confirmation loop before applying any fix')
bullet('Structured response options per check type (matching the existing Excel dropdown options)')
bullet('LLM parsing of free-text replies and generation of notification messages')
bullet('Agent writes confirmed corrections back to DHIS2 via the standard API')
bullet('Re-validation: after corrections are applied, re-run all checks on the corrected data')
bullet('Standalone Data Issue Log page showing all issues, conversations, and resolution history')
bullet('Configurable thresholds via a config file — no hardcoding')
spacer()

h2('Out of scope for MVP — Phase 2')
bullet('WhatsApp and email channels')
bullet('Pre-commit interception (requires a custom DHIS2 React app)')
bullet('Embedding the issue log as a native DHIS2 app')
bullet('Inbound phone call routing')
spacer()

# ── 9. Open questions for George ──────────────────────────────────────────────
h1('9. Open Questions')
bullet('Access to the existing R scripts — needed for Phase 2 to upgrade outlier detection to the 5-method ensemble; not a blocker for MVP')
bullet('Reporting deadline — what day of the following month triggers the missing report flag?')
bullet('Contact registry — do facility workers have registered mobile numbers, or is there a shared facility phone? Who owns and maintains that list?')
bullet('Dismissed warnings — if a facility acknowledges a flag but submits the original value anyway, should that automatically escalate to the woreda?')
bullet('Post-commit vs. pre-commit — is post-commit notification (within minutes of submission) acceptable for the MVP, or is blocking the form before submission a hard requirement?')
bullet('SMS language — should messages go out in Amharic, English, or both?')
spacer()

doc.save('AHEAD_AI_MVP_Architecture.docx')
print('Saved: AHEAD_AI_MVP_Architecture.docx')
