#!/usr/bin/env python3
"""Generates AHEAD_AI_UX_Workflow.docx — UX and workflow specification.
Run: pip install python-docx && python3 generate_ux_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

def h1(t):
    p = doc.add_heading(t, 1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
    p.runs[0].font.size = Pt(15)
def h2(t):
    p = doc.add_heading(t, 2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
    p.runs[0].font.size = Pt(12)
def body(t):
    p = doc.add_paragraph(t)
    if p.runs: p.runs[0].font.size = Pt(10.5)
def bullet(t, lv=0):
    p = doc.add_paragraph(t, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (lv + 1))
def mono(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(t)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shd)
def tbl(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].text = h
        run = hc[i].paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'D6E4F0')
        hc[i]._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            if cells[ci].paragraphs[0].runs:
                cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
def sp(): doc.add_paragraph('')


# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('AHEAD AI Agent — UX & Workflow Specification', 0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
title.runs[0].font.size = Pt(18)
sub = doc.add_paragraph('SMS Notification Design, Escalation Logic, and Issue Log  |  UNICEF AHEAD × Gates Foundation AI Fellows  |  June 2026')
sub.runs[0].font.size = Pt(9); sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88); sub.runs[0].italic = True
sp()

# ── 1. User Roles ─────────────────────────────────────────────────────────────
h1('1. User Roles and Notification Scope')
body('Six roles exist in the AHEAD hierarchy. For MVP, only facility workers and woreda HMIS officers receive active SMS notifications. Zone, regional, and national staff receive escalation summaries only when issues remain unresolved at lower levels. The AHEAD/UNICEF team monitors everything passively via the issue log web page.')
sp()
tbl(
    ['Role', 'Notification channel', 'What they receive', 'MVP scope'],
    [
        ['Facility worker',
         'SMS (individual)',
         'One SMS per issue with a reference ID; confirmation loop; resolution confirmation',
         'Active — first point of contact for every issue'],
        ['Woreda HMIS officer',
         'SMS (summary)',
         'Batched summary of all unresolved issues in their woreda after 72h',
         'Active — escalation target'],
        ['Zone officer',
         'SMS (summary)',
         'Summary of unresolved issues across woredas after 1 week',
         'Active — escalation target'],
        ['Regional focal point',
         'SMS (summary)',
         'Summary of unresolved issues across zones after 1 further week',
         'Active — escalation target'],
        ['National TWG',
         'SMS (monthly digest)',
         'End-of-month summary of all unresolved issues across Ethiopia',
         'Active — monthly digest only'],
        ['UNICEF / AHEAD team',
         'Issue log web page',
         'Real-time view of all issues, statuses, and full SMS threads',
         'Passive — observe and intervene via log page'],
    ]
)
sp()
body('Important distinction: the notification contacts (phone numbers) are stored in the agent\'s own contact registry, not in DHIS2 user accounts. DHIS2 user accounts control data-entry access. The contact registry controls who gets notified. These are maintained separately.')
sp()
body('Demo focus: The live demonstration centres entirely on the facility-level health worker — the person who receives the initial SMS, engages in the correction conversation, and sees their submission corrected in DHIS2. The escalation cascade (woreda → zone → region → national) is part of the production architecture and fires automatically in the background, but is not walked through in the demo. The AHEAD team monitors everything via the standalone issue log web app shown at the end of the demo.')
sp()

# ── 2. Notification Triggers ──────────────────────────────────────────────────
h1('2. Notification Triggers')
tbl(
    ['Issue type', 'What causes it', 'Detection timing', 'Who is notified first'],
    [
        ['Statistical outlier',
         'A submitted value is flagged by DHIS2 outlier detection (Z-score >= 3.0 vs. facility historical baseline)',
         'Within ~30 seconds of form submission',
         'Facility worker'],
        ['DTP1/DTP3 inconsistency',
         'Penta3 doses exceed Penta1 doses for same facility-period (biologically implausible)',
         'Within ~30 seconds of form submission',
         'Facility worker'],
        ['Missing report',
         'Facility has no completed dataset registration by day 5 of the following month',
         'Daily at 8am, starting day 5 of following month',
         'Facility worker'],
    ]
)
sp()

# ── 3. Message Templates ──────────────────────────────────────────────────────
h1('3. SMS Message Templates')
body('All outbound notifications present numbered multiple-choice options — the same response categories used in the AHEAD Excel dropdown workflow. The worker replies with a number (1, 2, 3 etc.). Options requiring additional input (a specific value, or a brief explanation) trigger a short follow-up question. The numbered selection itself is the decision — there is no separate YES/NO confirmation step.')
sp()

h2('3.1  Outlier — initial notification (6 options)')
body('Sent when a data value is flagged by outlier detection. Options 2 and 4 require a follow-up reply.')
mono(
    '[DQ-AX42] Addi Arekay HC, Jun 2026\n'
    'BCG under-1: 970 (normal ~97, flagged)\n'
    'Reply with option number:\n'
    '1. Replace with 6-month avg\n'
    '2. Keep as-is (explain)\n'
    '3. Treat as zero\n'
    '4. Replace with specific number\n'
    '5. Use facility doses only\n'
    '6. Use outreach doses only'
)
sp()

h2('3.2  Outlier — follow-up for option 2 (keep as-is)')
mono(
    '[DQ-AX42] Option 2: keep as-is.\n'
    'Please reply with a brief explanation.\n'
    '(e.g., "PIRI campaign in May" or "outreach week")'
)
sp()

h2('3.3  Outlier — follow-up for option 4 (specific number)')
mono(
    '[DQ-AX42] Option 4 selected.\n'
    'Reply with the correct BCG under-1 value.'
)
sp()

h2('3.4  DTP1/DTP3 inconsistency — initial notification (5 options)')
body('Sent when Penta3 > Penta1. Relative difference is shown so the worker can immediately see the scale of the gap. Options 1, 4, and 5 require a follow-up reply.')
mono(
    '[DQ-BK19] Addi Arekay HC, Jun 2026\n'
    'Penta1 under-1: 60 | Penta3 under-1: 90\n'
    'Penta3 exceeds Penta1 by 50%.\n'
    'Reply with option number:\n'
    '1. Keep as-is (explain)\n'
    '2. Set both to Penta1 value (60)\n'
    '3. Set both to Penta3 value (90)\n'
    '4. Replace with specific values\n'
    '5. Other (explain)'
)
sp()

h2('3.5  DTP — follow-up for options 1 and 5 (explain)')
mono(
    '[DQ-BK19] Option 1: keep as-is.\n'
    'Please reply with a brief explanation.'
)
sp()

h2('3.6  DTP — follow-up for option 4 (specific values)')
mono(
    '[DQ-BK19] Option 4 selected.\n'
    'Reply with correct Penta1 and Penta3 under-1\n'
    'separated by comma. (e.g., 45,42)'
)
sp()

h2('3.7  Missing report — initial notification (SUBMIT + 4 options)')
body('Recovery before imputation: the SUBMIT option is presented first so the worker can locate and submit the missing report. Cleaning options are a fallback only if recovery fails. This is explicit in the AHEAD methodology: "the first step is always recovery, not imputation."')
mono(
    '[DQ-MN33] Addi Arekay HC\n'
    'No completed EPI report for May 2026.\n'
    'Can you submit the report? Reply SUBMIT.\n'
    'Or if unavailable, select:\n'
    '1. Replace with 6-month avg\n'
    '2. Replace with zero (closed/no service)\n'
    '3. Replace with specific values\n'
    '4. Remove from analysis'
)
sp()

h2('3.8  Missing report — follow-up for SUBMIT')
mono(
    '[DQ-MN33] Thanks — checking DHIS2 within the hour.\n'
    'Issue will close automatically once your\n'
    'submission appears.'
)
sp()

h2('3.9  Missing report — follow-up for option 3 (specific values)')
mono(
    '[DQ-MN33] Option 3 selected.\n'
    'Reply with BCG, Penta1, Penta3, MR1 under-1\n'
    'separated by commas. (e.g., 38,34,28,33)'
)
sp()

h2('3.10  Resolution confirmation')
body('Sent after any option is applied. No YES/NO step — the numbered reply is the decision.')
mono(
    '[DQ-AX42] Done.\n'
    'BCG under-1 at Addi Arekay HC updated to\n'
    '6-month avg (96). Decision logged. Issue closed.'
)
sp()

h2('3.11  Retry (24h no reply — up to 3 times)')
mono(
    '[DQ-AX42] Reminder (2 of 3):\n'
    'Addi Arekay HC, Jun 2026 — BCG under-1 = 970\n'
    'Reply 1-6 to resolve. Escalates if no reply.'
)
sp()

h2('3.12  Woreda escalation summary')
body('Sent to the woreda HMIS officer after 72h of no facility-level resolution. One SMS per woreda covers all open issues in that woreda.')
mono(
    '[DQ-WOREDA] Addi Arekay Woreda\n'
    '2 unresolved issues, Jun 2026:\n'
    'DQ-AX42 Addi Arekay HC: BCG outlier (3 days)\n'
    'DQ-BK19 Addi Arekay HC: Penta3 > Penta1 (3 days)\n'
    'Details: <issue log URL>'
)
sp()

h2('3.13  Monthly national digest')
body('Sent to the national TWG on the first of each month covering all prior-month unresolved issues.')
mono(
    '[DQ-MONTHLY] Jun 2026 Summary\n'
    'Ethiopia: 2 issues still open\n'
    'Amhara / N. Gondar:\n'
    '  - 1 outlier (Addi Arekay HC, 25 days)\n'
    '  - 1 DTP inconsistency (Addi Arekay HC)\n'
    'View all: <issue log URL>'
)
sp()

# ── 4. Conversation Flows ─────────────────────────────────────────────────────
h1('4. Conversation Flows')
body('The numbered-option model keeps conversations short: notification → number selection → (optional follow-up for extra input) → resolution confirmation. Two to four turns total. The numbered reply IS the decision — no separate YES/NO confirmation.')
sp()

h2('4.1  Outlier — option with no follow-up (option 1: 6-month average)')
body('Agent calculates the 6-month average and applies it automatically. Two turns.')
mono(
    'Agent     [DQ-AX42] Addi Arekay HC, Jun 2026. BCG under-1: 970 (normal ~97).\n'
    '          Reply 1-6: 1=6mo-avg 2=keep(explain) 3=zero 4=specific# 5=HF-only 6=outreach-only\n'
    '\n'
    'Facility  1\n'
    '\n'
    'Agent     [DQ-AX42] Done. BCG under-1 replaced with 6-month avg (96).\n'
    '          Decision logged. Issue closed.'
)
sp()

h2('4.2  Outlier — option with follow-up (option 4: specific number)')
body('Three turns: notification, number selection, value entry.')
mono(
    'Agent     [DQ-AX42] Addi Arekay HC, Jun 2026. BCG under-1: 970 (normal ~97).\n'
    '          Reply 1-6: 1=6mo-avg 2=keep(explain) 3=zero 4=specific# 5=HF-only 6=outreach-only\n'
    '\n'
    'Facility  4\n'
    '\n'
    'Agent     [DQ-AX42] Option 4: replace with specific number.\n'
    '          Reply with the correct BCG under-1 value.\n'
    '\n'
    'Facility  35\n'
    '\n'
    'Agent     [DQ-AX42] Done. BCG under-1 updated to 97 at Addi Arekay HC.\n'
    '          Decision logged. Issue closed.'
)
sp()

h2('4.3  Outlier — keep as-is with explanation (option 2)')
body('Three turns: notification, option selection, explanation text.')
mono(
    'Agent     [DQ-AX42] Addi Arekay HC, Jun 2026. BCG under-1: 970 (normal ~97).\n'
    '          Reply 1-6: 1=6mo-avg 2=keep(explain) 3=zero 4=specific# 5=HF-only 6=outreach-only\n'
    '\n'
    'Facility  2\n'
    '\n'
    'Agent     [DQ-AX42] Option 2: keep as-is.\n'
    '          Please reply with a brief explanation.\n'
    '\n'
    'Facility  PIRI campaign in North Gondar, Jun 2026\n'
    '\n'
    'Agent     [DQ-AX42] Noted. BCG under-1 kept as 350.\n'
    '          Comment logged: "PIRI campaign in North Gondar, Jun 2026".\n'
    '          Issue closed.'
)
sp()

h2('4.4  DTP inconsistency — simple option (option 2: use Penta1 for both)')
body('Agent replaces both values immediately. Two turns.')
mono(
    'Agent     [DQ-BK19] Addi Arekay HC, Jun 2026.\n'
    '          Penta1: 60 | Penta3: 90 (Penta3 exceeds Penta1 by 50%).\n'
    '          Reply 1-5: 1=keep(explain) 2=use-Penta1(45) 3=use-Penta3(80) 4=specific 5=other\n'
    '\n'
    'Facility  2\n'
    '\n'
    'Agent     [DQ-BK19] Done. Penta1 and Penta3 under-1 both set to 60.\n'
    '          Decision logged. Issue closed.'
)
sp()

h2('4.5  DTP inconsistency — specific values (option 4)')
body('Three turns: notification, option selection, two values separated by comma.')
mono(
    'Agent     [DQ-BK19] Addi Arekay HC, Jun 2026.\n'
    '          Penta1: 60 | Penta3: 90 (Penta3 exceeds Penta1 by 50%).\n'
    '          Reply 1-5: 1=keep(explain) 2=use-Penta1(45) 3=use-Penta3(80) 4=specific 5=other\n'
    '\n'
    'Facility  4\n'
    '\n'
    'Agent     [DQ-BK19] Option 4: replace with specific values.\n'
    '          Reply with correct Penta1 and Penta3 under-1 separated by comma.\n'
    '\n'
    'Facility  88,72\n'
    '\n'
    'Agent     [DQ-BK19] Done. Penta1 under-1 = 45, Penta3 under-1 = 38.\n'
    '          Decision logged. Issue closed.'
)
sp()

h2('4.6  Missing report — recovery succeeds (SUBMIT)')
body('Worker submits the missing report in DHIS2. Agent re-checks within the hour and closes automatically.')
mono(
    'Agent     [DQ-MN33] Addi Arekay HC. No completed EPI report for May 2026.\n'
    '          Reply SUBMIT if submitting now, or 1-4 for cleaning.\n'
    '\n'
    'Facility  SUBMIT\n'
    '\n'
    'Agent     [DQ-MN33] Thanks — checking DHIS2 within the hour.\n'
    '          Issue will close automatically once your submission appears.\n'
    '\n'
    '          [1 hour later — agent finds the registration]\n'
    '\n'
    'Agent     [DQ-MN33] Confirmed. Addi Arekay HC May 2026 report received. Issue closed.'
)
sp()

h2('4.7  Missing report — cleaning option (option 2: replace with zero)')
body('Recovery not possible. Worker selects a cleaning option instead.')
mono(
    'Agent     [DQ-MN33] Addi Arekay HC. No completed EPI report for May 2026.\n'
    '          Reply SUBMIT if submitting now, or 1-4 for cleaning.\n'
    '\n'
    'Facility  2\n'
    '\n'
    'Agent     [DQ-MN33] Done. May 2026 recorded as zero for all vaccines at Addi Arekay HC.\n'
    '          (Facility closed / no service that month.) Issue closed.'
)
sp()

h2('4.8  No reply — retry then escalation')
mono(
    'Day 0     Agent sends numbered-option notification\n'
    'Day 1     No reply — Retry 1: "[DQ-AX42] Reminder (1 of 3): Addi Arekay HC BCG=970. Reply 1-6."\n'
    'Day 2     No reply — Retry 2: "[DQ-AX42] Reminder (2 of 3): ... Issue escalates if no reply."\n'
    'Day 3     No reply — Retry 3 + escalate: woreda HMIS officer receives summary SMS\n'
    'Day 10    No woreda resolution — escalate to zone\n'
    'Day 17    No zone resolution — escalate to region\n'
    'Month end National digest includes all unresolved issues'
)
sp()

# ── 5. Escalation Timeline ────────────────────────────────────────────────────
h1('5. Escalation Timeline')
body('The timeline below shows a single issue from detection to national digest, assuming no resolution at any level.')
sp()
mono(
    'Day 0    Issue detected → facility worker notified (SMS)\n'
    '         ├── Facility replies → confirmation loop → resolved (ideal path)\n'
    '         └── No reply...\n'
    '\n'
    'Day 1    Retry 1 of 3 (24h after initial)\n'
    'Day 2    Retry 2 of 3 (48h after initial)\n'
    'Day 3    Retry 3 of 3 + escalate to woreda HMIS officer\n'
    '         ├── Woreda intervenes → resolved\n'
    '         └── No resolution...\n'
    '\n'
    'Day 10   Escalate to zone officer (1 week after woreda notification)\n'
    '         ├── Zone intervenes → resolved\n'
    '         └── No resolution...\n'
    '\n'
    'Day 17   Escalate to regional focal point (1 week after zone notification)\n'
    '         └── No resolution...\n'
    '\n'
    'Month end  National TWG receives monthly digest (includes this issue)'
)
sp()
body('Escalation messages at woreda level and above are sent as a single summary SMS covering all unresolved issues from that org unit in the same period — not one SMS per issue. This prevents notification fatigue at higher levels where one officer may receive reports from many facilities.')
sp()

# ── 6. Issue Log Web App ──────────────────────────────────────────────────────
h1('6. Issue Log Web App (Standalone)')
body('A standalone web application at http://localhost:5001, completely separate from DHIS2. No DHIS2 login required. It reads from the agent\'s own SQLite database and is the primary monitoring surface for the AHEAD team and supervisors. Every issue the agent detects, every SMS sent and received, and every state change is visible here in real-time.')
sp()

h2('Page layout (wireframe)')
mono(
    'AHEAD Data Issue Log                      [ Open: 2 | Resolved: 47 | All: 49 ]\n'
    'Filter: [ All types v ]  [ All woredas v ]  [ Apr 2026 v ]  [ Search ref ID... ]\n'
    '--------------------------------------------------------------------------------\n'
    'Ref     Facility          Issue              Period  Status    Opened      Updated\n'
    '--------------------------------------------------------------------------------\n'
    'DQ-AX42 Addi Arekay HC    Outlier: BCG 970   202606  RESOLVED  2026-06-01  06-01\n'
    'DQ-BK19 Addi Arekay HC    DTP3 > DTP1        202606  RESOLVED  2026-06-01  06-01\n'
    'DQ-MN33 Addi Arekay HC    Missing report     202605  RESOLVED  2026-06-01  06-01\n'
    '[> expand] Full SMS thread for DQ-AX42:\n'
    '  2026-06-01 10:02  OUT  [DQ-AX42] Addi Arekay HC, Jun 2026. BCG under-1: 970...\n'
    '  2026-06-01 10:04  IN   4\n'
    '  2026-06-01 10:04  OUT  [DQ-AX42] Option 4: replace with specific number...\n'
    '  2026-06-01 10:05  IN   97\n'
    '  2026-06-01 10:05  OUT  [DQ-AX42] Done. BCG under-1 updated to 97. Issue closed.\n'
    '--------------------------------------------------------------------------------'
)
sp()

h2('Issue log columns')
tbl(
    ['Column', 'Content'],
    [
        ['Ref ID',          'e.g. DQ-AX42 — unique across all issues, never reused'],
        ['Facility',        'Facility name and woreda in parentheses'],
        ['Issue type',      'Outlier (element + value), DTP3 > DTP1, or Missing report'],
        ['Period',          'DHIS2 period code, displayed as human-readable month (Apr 2026)'],
        ['Status',          'OPEN | AWAITING CONFIRM | RESOLVED | DISMISSED | ESCALATED'],
        ['Cascade level',   'Current level: Facility / Woreda / Zone / Region / National'],
        ['Opened',          'Timestamp when issue was first detected'],
        ['Last contact',    'Timestamp of last outbound SMS'],
        ['Resolved value',  'What the value was corrected to (blank if not yet resolved)'],
        ['[Expand]',        'Click to see the full SMS thread for that issue'],
    ]
)
sp()

h2('Filters')
tbl(
    ['Filter', 'Options'],
    [
        ['Status',   'All / Open / Escalated / Resolved / Dismissed'],
        ['Type',     'All / Outlier / DTP inconsistency / Missing report'],
        ['Woreda',   'All / [woreda names from org unit hierarchy]'],
        ['Period',   'All / [YYYYMM dropdown]'],
        ['Search',   'Free-text search on ref ID or facility name'],
    ]
)
sp()

# ── 7. Edge Cases ─────────────────────────────────────────────────────────────
h1('7. Edge Cases and Agent Behaviour')
tbl(
    ['Edge case', 'Agent behaviour'],
    [
        ['Inbound SMS contains no reference ID',
         'Agent checks if the sender\'s phone number has exactly one open issue. If yes, routes to that issue. If multiple open issues, replies: "Please include your issue ref ID (e.g. DQ-AX42) so we can route your reply."'],
        ['Reference ID not found or already closed',
         'Agent replies: "Ref ID not recognised or issue already closed. View all issues at <issue log URL>."'],
        ['Worker replies with non-numeric text for an outlier issue',
         'Claude returns intent=unknown. Agent replies: "Please reply with a number (the correct value) or KEEP to leave as-is."'],
        ['Worker sends a value that still looks like an outlier',
         'Agent applies the confirmed value as requested, then re-runs DQ checks in the next poll cycle. If the new value still fails, a new issue is created with a new ref ID.'],
        ['Same facility has two open issues simultaneously',
         'Both have distinct ref IDs. Every outbound message includes the ref ID. Worker must include the ref ID in replies — if missing, agent uses the single-open-issue routing logic above.'],
        ['Correction POST to DHIS2 fails (DHIS2 down or network error)',
         'Agent sends SMS: "Could not apply the change right now. Please re-enter the value directly in DHIS2 and reply DONE once submitted." Issue status remains awaiting_confirm. Retry flagged for next cycle.'],
        ['Worker replies DONE (missing report already submitted)',
         'Agent schedules a re-check in 1 hour via APScheduler. If the registration appears, issue closes automatically with a confirmation SMS. If not, issue remains open.'],
        ['Escalation notification itself fails to deliver',
         'Twilio delivery failure is logged. Agent retries the escalation SMS once after 30 minutes. Failure is visible on the issue log page under last_contact.'],
        ['Issue resolved at woreda level (not facility)',
         'If woreda officer contacts AHEAD team directly and they manually close the issue via the log page, status is set to resolved with a note. No further SMS is sent.'],
    ]
)
sp()

# ── 8. Reference ID Design ────────────────────────────────────────────────────
h1('8. Reference ID Design')
body('Every issue gets a unique ID on creation. The ID appears in every outbound message and is how the agent routes inbound replies to the right issue. Design constraints:')
sp()
bullet('Visually unambiguous when read aloud or off a small phone screen')
bullet('Short enough to quote in an SMS without taking up too much space')
bullet('Random enough to be hard to guess (prevents spoofed replies)')
bullet('Never reused — resolved issues keep their ID permanently for audit')
sp()
tbl(
    ['Property', 'Decision', 'Reason'],
    [
        ['Format', 'DQ-XXXX (4-character suffix)', 'DQ prefix identifies the source system; 4 chars is short and unambiguous'],
        ['Character set',
         'A B C D E F G H J K M N P Q R S T U V W X Y + 3 4 5 6 7 8 9 (29 chars)',
         'Excludes 0/O, 1/I/L, 2/Z — characters that look alike on paper or screen'],
        ['Capacity', '29^4 = 707,281 unique IDs', 'More than sufficient for MVP; Phase 2 can extend to 5 chars if needed'],
        ['Generation', 'Python secrets.choice() from the charset', 'Cryptographically random — hard to guess or enumerate'],
        ['Collision check', 'DB lookup before issuing', 'Probability of collision is < 0.001% even at 10,000 open issues'],
    ]
)
sp()

# ── 9. MVP vs Phase 2 Scope ───────────────────────────────────────────────────
h1('9. MVP vs Phase 2 Scope')
body('The following table clarifies what is in scope for this build and what is explicitly deferred. Phase 2 items are excluded not because they are unimportant but because they either require external dependencies (Meta WhatsApp approval, DHIS2 App Framework expertise) or can be validated after the core SMS loop is working.')
sp()
tbl(
    ['Feature', 'MVP', 'Phase 2', 'Notes'],
    [
        ['Post-commit DQ checks (polling)',                     'Yes', '—',  ''],
        ['Outlier detection (DHIS2 built-in, Z-score)',         'Yes', '—',  ''],
        ['DTP1/DTP3 consistency (DHIS2 validation rules)',      'Yes', '—',  ''],
        ['Missing report detection',                            'Yes', '—',  ''],
        ['Two-way SMS conversation with confirmation loop',     'Yes', '—',  ''],
        ['Full escalation cascade (facility → national)',       'Yes', '—',  ''],
        ['Apply corrections back to DHIS2 via API',            'Yes', '—',  ''],
        ['Issue log web page (standalone Flask)',               'Yes', '—',  ''],
        ['Claude API (reply parsing + message generation)',     'Yes', '—',  ''],
        ['English-language messages only',                      'Yes', '—',  ''],
        ['5-method outlier ensemble (R script integration)',    '—',  'Yes', 'Requires access to AHEAD R scripts; deferred'],
        ['Name consistency check',                              '—',  'Yes', 'Not covered by DHIS2 built-in rules'],
        ['WhatsApp channel',                                    '—',  'Yes', 'Requires Meta Business API approval (1-4 weeks)'],
        ['Email channel',                                       '—',  'Yes', 'Reply parsing noisier; lower priority'],
        ['Pre-commit inline form warning (custom form JS)',     '—',  'Yes', 'Requires DHIS2 custom form; useful but not needed to validate cascade'],
        ['Amharic messages',                                    '—',  'Yes', 'Claude can generate; need translation review before sending'],
        ['DHIS2 embedded app (issue log inside DHIS2)',        '—',  'Yes', 'Requires DHIS2 App Framework (React); standalone page sufficient for MVP'],
        ['Inbound phone call routing',                          '—',  'Yes', 'Higher complexity; SMS covers facility workers adequately for MVP'],
    ]
)
sp()

# ── 10. Demo Script ───────────────────────────────────────────────────────────
h1('10. Demo Script')
body('Three live data-entry demos, all using the same facility and user account: eth_facility_01 (Almaz Tadesse), Addi Arekay Health Center. You type wrong values directly into the DHIS2 form as the facility worker would, the agent detects the issue within ~30 seconds, and the numbered-option SMS conversation plays out on the presenter\'s phone. Total runtime: ~15 minutes.')
sp()

tbl(
    ['Demo', 'Issue type', 'Facility', 'User', 'Period', 'Duration', 'What it proves'],
    [
        ['A', 'Statistical outlier',     'Addi Arekay HC', 'eth_facility_01', 'Jun 2026', '~6 min', 'Type BCG=970 (meant 97) → outlier detected → option 4 → corrected to 97 in DHIS2'],
        ['B', 'DTP1/DTP3 inconsistency', 'Addi Arekay HC', 'eth_facility_01', 'Jun 2026', '~5 min', 'Edit Penta3=90 > Penta1=60 → validation rule → option 2 → both corrected'],
        ['C', 'Missing report',          'Addi Arekay HC', 'eth_facility_01', 'May 2026', '~4 min', 'No data for prior month → agent flags → SUBMIT → recovery flow'],
        ['D', 'Issue log web app',       '(all)',           '—',               '(all)',    '~3 min', 'Supervisor view: all issues, numbered-option threads, statuses'],
    ]
)
sp()

h2('Prerequisites — confirm before presenting')
bullet('DHIS2 running: http://localhost:8080 responds with login page')
bullet('Agent running: http://localhost:5001/api/status returns JSON with uptime')
bullet('inject_data.py has been run — 28 months of clean baseline data loaded (Jan 2024 – Apr 2026). No pre-seeded issues. Addi Arekay Health Center has 22 months of clean baseline values.')
bullet('contacts table in agent SQLite seeded with presenter\'s phone for Addi Arekay Health Center')
bullet('Presenter\'s phone visible to audience (screen mirror or camera pointing at phone)')
bullet('Twilio account active and TWILIO_FROM_NUMBER configured in .env')
sp()

body('All demos use a single login: eth_facility_01 / [DHIS2_USER_PASSWORD from .env]. This account belongs to Almaz Tadesse and is scoped to Addi Arekay Health Center only — exactly the facility-worker persona being demonstrated.')
sp()
body('Demos A and B use June 2026 (202606) — a period with no data yet. Demo C uses May 2026 (202505) to show the missing report scenario for a prior month that was never formally completed.')
sp()
body('Addi Arekay Health Center is a Health Center (tier 2). BCG under-1 baseline: ~97/month (range 80–115 across 22 months). Penta1 under-1: ~87. Penta3 under-1: ~71. Detection is within ~30 seconds via lastUpdated polling, or trigger instantly with: POST http://localhost:5001/api/scan')
sp()

h2('Demo A — Statistical Outlier: Addi Arekay Health Center, Jun 2026 (~6 min)')
body('What this shows: Almaz enters her June 2026 BCG count with an extra zero — a realistic typo. The agent detects it against 22 months of baseline history and sends the numbered-option SMS. The full loop plays out in front of the audience in real time.')
sp()

body('Step 1 — Log in and open the data entry form')
bullet('Log in at http://localhost:8080 as eth_facility_01 / [DHIS2_USER_PASSWORD from .env]')
bullet('Open Data Entry (grid icon in top menu)')
bullet('Navigate in the left org unit tree: Ethiopia → Amhara → North Gondar → Addi Arekay Woreda → Addi Arekay Health Center')
bullet('Select dataset: EPI - Routine vaccine delivery | Period: June 2026')
bullet('The form is blank — no data for June 2026 yet')
bullet('Narrate: "Almaz is entering the monthly EPI report for June. She works at the facility and submits this form each month."')
sp()

body('Step 2 — Enter data with a BCG outlier (deliberate typo)')
bullet('Enter the following values in the form:')
bullet('BCG under-1: 970  |  BCG >= 1 year: 5', lv=1)
bullet('Penta1 under-1: 88  |  Penta1 >= 1 year: 4', lv=1)
bullet('Penta3 under-1: 72  |  Penta3 >= 1 year: 3', lv=1)
bullet('MR1 under-1: 83  |  MR1 >= 1 year: 3', lv=1)
bullet('Narrate: "BCG under-1 should be 97 this month — but Almaz accidentally typed 970. An extra zero. This happens."')
bullet('Click Save')
sp()

body('Step 3 — Agent detects the outlier (~30 seconds)')
bullet('Wait ~30 seconds, OR trigger immediately: curl -s -X POST http://localhost:5001/api/scan')
bullet('Agent queries outlier detection: BCG = 970 vs 22-month mean ~97, z-score ~87 — clearly flagged')
bullet('New issue created. Open http://localhost:5001/issues — issue appears as OPEN')
sp()

body('Step 4 — SMS arrives on presenter\'s phone')
mono(
    '[DQ-XXXX] Addi Arekay HC, Jun 2026\n'
    'BCG under-1: 970 (normal ~97, flagged)\n'
    'Reply with option number:\n'
    '1. Replace with 6-month avg\n'
    '2. Keep as-is (explain)\n'
    '3. Treat as zero\n'
    '4. Replace with specific number\n'
    '5. Use facility doses only\n'
    '6. Use outreach doses only'
)
bullet('Point out: numbered options mirror the AHEAD Excel dropdown exactly — same decision schema George\'s team uses today, now delivered by SMS')
sp()

body('Step 5 — Select option 4 (replace with specific number)')
bullet('Reply to the SMS: 4')
bullet('Agent sends a follow-up:')
mono(
    '[DQ-XXXX] Option 4: replace with specific number.\n'
    'Reply with the correct BCG under-1 value.'
)
bullet('Reply: 97')
bullet('Agent applies the correction and sends:')
mono(
    '[DQ-XXXX] Done. BCG under-1 updated to 97 at Addi Arekay HC.\n'
    'Decision logged. Issue closed.'
)
sp()

body('Step 6 — Show the correction in DHIS2')
bullet('Switch back to DHIS2 and refresh Addi Arekay HC / June 2026')
bullet('BCG under-1 now shows 97 (was 970)')
bullet('Narrate: "The correction was written to DHIS2 by the agent. The audit log will show agent_service made this change — fully traceable."')
sp()

body('Expected final state:')
bullet('DHIS2: Addi Arekay HC, Jun 2026, BCG under-1 = 97 (was 970)')
bullet('Issue log: DQ-XXXX | Outlier: BCG | Addi Arekay HC | RESOLVED | Option 4: specific number (97)')
bullet('SMS thread: 4 messages — notification, "4", follow-up question, "97", resolution')
sp()

h2('Demo B — DTP1/DTP3 Inconsistency: Addi Arekay Health Center, Jun 2026 (~5 min)')
body('What this shows: Almaz accidentally swaps her Penta1 and Penta3 figures. The EPI metadata package validation rule fires — Penta3 > Penta1 by 50%, well above the 30% threshold for monthly facility data. Same form, same login.')
sp()

body('Step 1 — Edit the same June 2026 form')
bullet('Still logged in as eth_facility_01. Navigate back to Addi Arekay HC / June 2026.')
bullet('BCG is now 97 (corrected in Demo A). The Penta values currently read: Penta1 under-1: 88, Penta3 under-1: 72.')
bullet('Update: Penta1 under-1 = 60, Penta3 under-1 = 90')
bullet('Narrate: "Almaz has swapped the Penta1 and Penta3 figures — maybe copied from the wrong row on her tally sheet."')
bullet('Click Save')
sp()

body('Step 2 — Agent detects the DTP violation (~30 seconds)')
bullet('Wait or POST http://localhost:5001/api/scan')
bullet('Agent runs the DHIS2 validation check. Built-in EPI rule fires: Penta3 (90) > Penta1 (60)')
bullet('Relative difference: (90-60)/60 = 50% — above the 30% threshold for monthly facility data')
bullet('New issue created and visible in http://localhost:5001/issues')
sp()

body('Step 3 — SMS arrives')
mono(
    '[DQ-XXXX] Addi Arekay HC, Jun 2026\n'
    'Penta1 under-1: 60 | Penta3 under-1: 90\n'
    'Penta3 exceeds Penta1 by 50%.\n'
    'Reply with option number:\n'
    '1. Keep as-is (explain)\n'
    '2. Set both to Penta1 value (60)\n'
    '3. Set both to Penta3 value (90)\n'
    '4. Replace with specific values\n'
    '5. Other (explain)'
)
bullet('Narrate: "The 3rd dose cannot exceed the 1st dose — dropout only goes one direction. The agent surfaces the gap (50%) and gives Almaz the same options she\'d see in the Excel sheet."')
sp()

body('Step 4 — Select option 4 (replace with specific values) to restore the correct numbers')
bullet('Narrate: "Almaz knows the real numbers — Penta1 is 88 and Penta3 is 72. She selects option 4 to enter both."')
bullet('Reply: 4')
bullet('Agent sends:')
mono(
    '[DQ-XXXX] Option 4: replace with specific values.\n'
    'Reply with correct Penta1 and Penta3 under-1\n'
    'separated by comma. (e.g., 45,42)'
)
bullet('Reply: 88,72')
bullet('Agent applies both and sends:')
mono(
    '[DQ-XXXX] Done. Penta1 under-1 = 88, Penta3 under-1 = 72.\n'
    'Decision logged. Issue closed.'
)
sp()

body('Step 5 — Show in DHIS2 and issue log')
bullet('Refresh Addi Arekay HC / June 2026: Penta1 under-1 = 88, Penta3 under-1 = 72')
bullet('Issue log: DQ-XXXX | DTP3 > DTP1 | Addi Arekay HC | RESOLVED | Option 4: specific values (88, 72)')
sp()

body('Expected final state:')
bullet('DHIS2: Addi Arekay HC, Jun 2026, Penta1=88, Penta3=72 — both corrected')
bullet('SMS thread: 4 messages — notification, "4", follow-up question, "88,72", resolution')
sp()

h2('Demo C — Missing Report: Addi Arekay Health Center, May 2026 (~4 min)')
body('What this shows: The agent proactively flags a prior month (May 2026) where no complete dataset registration exists for Addi Arekay HC. This is a different scenario from Demos A and B — rather than catching errors at entry time, it catches a facility that never submitted at all. The AHEAD methodology is explicit: recovery first, imputation only as a fallback.')
sp()

body('Step 1 — Show that May 2026 has no submission')
bullet('Navigate to Addi Arekay HC | Dataset: EPI - Routine vaccine delivery | Period: May 2026')
bullet('The form has data values (loaded by inject_data.py) but no complete registration — the "Complete" button was never clicked')
bullet('Narrate: "In the real workflow, the agent checks which facilities have submitted and formally completed their report. Addi Arekay HC has data values for May 2026 but has never been formally completed — the HMIS officer never clicked the Complete button. The agent flags this."')
sp()

body('Step 2 — Trigger the missing report check')
bullet('POST http://localhost:5001/api/scan?check=missing_reports')
bullet('Agent queries completeDataSetRegistrations for May 2026; Addi Arekay HC has no registration')
bullet('Issue created: type=missing_report, facility=Addi Arekay HC, period=May 2026')
sp()

body('Step 3 — SMS arrives on presenter\'s phone')
mono(
    '[DQ-XXXX] Addi Arekay HC\n'
    'No completed EPI report for May 2026.\n'
    'Can you submit the report? Reply SUBMIT.\n'
    'Or if unavailable, select:\n'
    '1. Replace with 6-month avg\n'
    '2. Replace with zero (closed/no service)\n'
    '3. Replace with specific values\n'
    '4. Remove from analysis'
)
bullet('Point out: SUBMIT appears first — the AHEAD methodology says recovery beats imputation; only choose a cleaning option if the report genuinely cannot be found')
sp()

body('Step 4 — Reply SUBMIT')
bullet('Reply: SUBMIT')
bullet('Agent sends:')
mono(
    '[DQ-XXXX] Thanks — checking DHIS2 within the hour.\n'
    'Issue will close automatically once your\n'
    'report is marked complete.'
)
bullet('Now in DHIS2: navigate to Addi Arekay HC / May 2026, click the Complete button to mark it as formally submitted')
bullet('Trigger a re-check: POST http://localhost:5001/api/scan?check=missing_reports')
bullet('Agent finds the registration and sends:')
mono(
    '[DQ-XXXX] Confirmed. Addi Arekay HC May 2026\n'
    'report received. Issue closed.'
)
sp()

body('Expected final state:')
bullet('DHIS2: Addi Arekay HC, May 2026 — marked as complete (Complete button clicked)')
bullet('Issue log: DQ-XXXX | Missing report | Addi Arekay HC | RESOLVED')
bullet('SMS thread: 3 messages — notification, "SUBMIT", acknowledgement, auto-close confirmation')
sp()

h2('Demo D — Issue Log Web App (~3 min)')
body('What this shows: The standalone supervisor view. George\'s team uses this to monitor all DQ activity across Ethiopia — every issue the agent detected, every SMS sent, every decision made — without ever logging into DHIS2.')
sp()

body('Step 1 — Open the issue log')
bullet('Navigate to http://localhost:5001/issues in the browser')
bullet('Three issues visible: Demo A (RESOLVED, BCG outlier), Demo B (RESOLVED, DTP inconsistency), Demo C (RESOLVED, missing report)')
sp()

body('Step 2 — Show filtering')
bullet('Filter by Status = "Open" — shows only unresolved issues (none after all three demos resolved)')
bullet('Filter by Facility = "Addi Arekay HC" — shows all three issues for this facility')
bullet('Narrate: "In production with hundreds of facilities across Ethiopia, George\'s team could filter by any woreda, region, or time period."')
sp()

body('Step 3 — Expand Demo A conversation thread (BCG outlier)')
bullet('Click expand on the BCG outlier issue (DQ-XXXX)')
bullet('Show the full SMS thread: timestamps, IN/OUT, each turn')
bullet('The thread shows: initial notification → "4" (option selected) → follow-up question → "97" (value given) → resolution')
bullet('Narrate: "This is the audit trail. It captures the same information the Excel dropdown captured — which option was chosen, what the corrected value is — but in real time, with the full conversation history."')
sp()

body('Step 4 — Show the decision schema link')
bullet('Point out: the issue record shows "Option 4: specific number (97)" — the equivalent of an Excel dropdown selection')
bullet('Narrate: "Every decision maps directly to the AHEAD response schema that George\'s team already uses. The cleaned dataset output would come from these logged decisions, exactly like it does today from the Excel files."')
sp()

body('Expected final state:')
bullet('Audience has seen the complete loop: facility entry → real-time detection → numbered SMS options → correction applied → audit log with full thread')
bullet('Issue log shows 3 resolved issues for Addi Arekay HC across two periods, with decision records')
sp()

h2('Common demo questions and answers')
tbl(
    ['Question', 'Answer'],
    [
        ['Why numbered options instead of free text?',
         'These are the same fixed categories in the current AHEAD Excel dropdowns. Using them means the agent captures exactly the same decisions the manual process captures, in a format the AHEAD team already understands.'],
        ['What if the facility worker ignores the SMS?',
         'The agent retries every 24h for 3 days. After 72h with no response, it automatically sends a summary to the woreda HMIS officer. No manual escalation needed.'],
        ['Who maintains the phone number list?',
         'The AHEAD team maintains a contacts registry (a simple table mapping each org unit to a phone number). It is loaded into the agent at startup and updated as staff change.'],
        ['Does the facility worker need to install an app?',
         'No. Standard SMS only. Works on any phone, including basic feature phones without internet.'],
        ['Can the agent send in Amharic?',
         'Not in MVP. Claude can generate Amharic and this is a Phase 2 upgrade pending translation review.'],
        ['What if the agent applies the wrong value?',
         'The agent logs every numbered option selection and the exact value applied. The issue thread is a complete audit trail. If a mistake is made, the AHEAD team can see it in the log and manually correct in DHIS2.'],
        ['What happens to issues that never get resolved?',
         'They escalate automatically through the hierarchy (facility → woreda → zone → region → national digest) and remain open in the log indefinitely until closed.'],
    ]
)
sp()

doc.save('AHEAD_AI_UX_Workflow.docx')
print('Saved: AHEAD_AI_UX_Workflow.docx')
