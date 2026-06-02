#!/usr/bin/env python3
"""Generates AHEAD_AI_Tech_Architecture.docx — MVP implementation specification.
Run: pip install python-docx && python3 generate_tech_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

def h1(t):
    p = doc.add_heading(t, 1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
    p.runs[0].font.size = Pt(15)
def h2(t):
    p = doc.add_heading(t, 2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
    p.runs[0].font.size = Pt(12)
def body(t):
    p = doc.add_paragraph(t)
    if p.runs: p.runs[0].font.size = Pt(10.5)
def bullet(t, lv=0):
    p = doc.add_paragraph(t, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (lv + 1))
def mono(t):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(t)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shd)
def tbl(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].text = h
        run = hc[i].paragraphs[0].runs[0]; run.bold = True; run.font.size = Pt(10)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'D6E4F0')
        hc[i]._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            if cells[ci].paragraphs[0].runs:
                cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
def sp(): doc.add_paragraph('')


# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('AHEAD AI Agent — Technical Architecture', 0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
title.runs[0].font.size = Pt(18)
sub = doc.add_paragraph('MVP Implementation Specification  |  UNICEF AHEAD × Gates Foundation AI Fellows  |  June 2026')
sub.runs[0].font.size = Pt(9); sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88); sub.runs[0].italic = True
sp()

# ── 1. Overview ───────────────────────────────────────────────────────────────
h1('1. Overview')
body('This document specifies the technical implementation of the AHEAD AI agent MVP. The agent monitors immunization data quality across the Ethiopia org unit hierarchy by running three automated checks against the DHIS2 REST API after each data submission. When an issue is detected, it drives a structured SMS cascade — notifying the responsible person at the facility level first, retrying for 72 hours, then escalating upward through woreda, zone, region, and national — until the issue is resolved or dismissed. The agent runs as a standalone Python service alongside the existing DHIS2 Docker instance, requiring no changes to DHIS2 itself.')
sp()
body('This specification covers the DHIS2 service account, DQ engine checks (with exact API endpoints), SQLite database schema, cascade state machine, Claude API integration, Twilio SMS integration, Flask application endpoints, APScheduler jobs, Docker Compose deployment, and all environment variables.')
sp()
body('Historical data note: the DHIS2 instance is loaded with 28 months of synthetic EPI data (Jan 2024 – Apr 2026) across all 12 facilities. The 22-month baseline (Jan 2024 – Oct 2025) provides the historical distribution required for meaningful Z-score outlier detection. The 6-month active window (Nov 2025 – Apr 2026) contains the three seeded DQ issues in the final period.')
sp()

# ── 2. Architecture Diagram ───────────────────────────────────────────────────
h1('2. Architecture Diagram')
mono(
    '  Docker host\n'
    '  ┌─────────────────────────────────────────────────────────────────────┐\n'
    '  │                                                                     │\n'
    '  │  ┌──────────────┐    REST API     ┌───────────────────────────────┐ │\n'
    '  │  │  DHIS2 2.40  │◄───────────────│  AHEAD Agent (Python)         │ │\n'
    '  │  │  port 8080   │                │                               │ │\n'
    '  │  └──────────────┘                │  ┌─────────────┐              │ │\n'
    '  │                                  │  │  DQ Engine  │              │ │\n'
    '  │  ┌──────────────┐                │  │  3 checks   │              │ │\n'
    '  │  │  PostgreSQL  │  (DHIS2 only)  │  └──────┬──────┘              │ │\n'
    '  │  └──────────────┘                │         │ issues              │ │\n'
    '  │                                  │  ┌──────▼──────┐              │ │\n'
    '  │  ┌──────────────┐                │  │State Machine│              │ │\n'
    '  │  │  SQLite DB   │◄───────────────│  │  cascade    │              │ │\n'
    '  │  │  agent/db/   │                │  └─────────────┘              │ │\n'
    '  │  └──────────────┘                │                               │ │\n'
    '  │                                  │  ┌─────────────┐ ┌──────────┐ │ │\n'
    '  │                                  │  │    Flask    │ │APScheduler│ │ │\n'
    '  │                                  │  │  port 5001  │ │  jobs    │ │ │\n'
    '  │                                  │  └─────────────┘ └──────────┘ │ │\n'
    '  │                                  └───────────────────────────────┘ │\n'
    '  └─────────────────────────────────────────────────────────────────────┘\n'
    '\n'
    '  Inbound SMS  ──►  Twilio  ──►  POST /webhook/sms  ──►  state machine\n'
    '  Outbound SMS ◄──  Twilio  ◄──  agent sends via Twilio REST API\n'
    '  Claude API   ◄──► agent (parse replies, generate messages)\n'
    '  Supervisors  ──►  browser  ──►  GET /issues  (issue log page)'
)
sp()

# ── 3. DHIS2 Service Account ──────────────────────────────────────────────────
h1('3. DHIS2 Service Account')
body('The agent authenticates to DHIS2 as a dedicated service account (agent_service) with national-level scope. This is separate from all human role-based accounts. Human accounts enforce data-entry access control; the agent account exists solely to read all data for DQ checks and write confirmed corrections. It must never be used for manual data entry.')
sp()
tbl(
    ['Property', 'Value', 'Reason'],
    [
        ['Username', 'agent_service',
         'Identifiable in DHIS2 audit log; distinct from human accounts'],
        ['Data entry org unit', 'Ethiopia (national root)',
         'Must write corrections at any facility in the hierarchy'],
        ['Data view org unit', 'Ethiopia (national root)',
         'Must read data from any facility to run DQ checks'],
        ['User role authorities', 'F_DATAVALUE_ADD, F_DATAVALUE_DELETE',
         'Read + write dataValueSets only — no metadata, settings, or user management'],
        ['Created by', 'build_users.py (add agent block)',
         'Reproducible setup; credentials set via AGENT_USER / AGENT_PASS in .env'],
    ]
)
sp()
body('All DHIS2 API calls made by the agent use HTTP Basic Auth with these credentials, read from environment variables at startup. Credentials are never hardcoded.')
sp()

# ── 4. DQ Engine ──────────────────────────────────────────────────────────────
h1('4. DQ Engine: Three Checks')
body('The agent detects new submissions by polling DHIS2\'s lastUpdated parameter every 30 seconds — a lightweight query that costs essentially nothing when nothing has changed. Only when changes are found does the agent run targeted DQ checks against the specific (orgUnit, period) pairs that have new data. Missing reports are the exception: detecting an absence cannot be event-driven, so they run on a separate daily cron.')
sp()

h2('4.0  Change detection (lastUpdated poll)')
body('Every 30 seconds the agent queries which data values have changed since the previous check. The response is an empty list when nothing was submitted — the DQ checks below are never triggered in that case. When changes are present, the agent extracts the affected (orgUnit, period) pairs and passes them to the DQ checks as a targeted scope, rather than scanning all facilities.')
mono(
    'GET /api/dataValueSets\n'
    '  ?dataSet=vI4ihClxSm4\n'
    '  &orgUnit=RFhqluFmvRG\n'
    '  &lastUpdated={iso_timestamp_of_last_check}\n'
    '  &fields=orgUnit,period\n'
    '\n'
    'If response.dataValues is empty: do nothing, update last_checked timestamp\n'
    'If non-empty: extract unique (orgUnit, period) pairs\n'
    '             pass to outlier + DTP checks as scoped targets\n'
    '             update last_checked timestamp'
)
sp()

h2('4.1  Outlier detection')
body('Uses the DHIS2 built-in outlier detection endpoint, which computes Z-scores for each data value against the same facility\'s historical baseline. The threshold defaults to 3.0 standard deviations and is configurable in config.py without a code change.')
mono(
    'GET /api/outlierDetection\n'
    '  ?ds=vI4ihClxSm4          # EPI - Routine vaccine delivery\n'
    '  &ou=RFhqluFmvRG          # Ethiopia root — scans all children\n'
    '  &startDate=2025-11-01\n'
    '  &endDate=2026-04-30\n'
    '  &algorithm=Z_SCORE\n'
    '  &threshold=3.0\n'
    '  &maxResults=100\n'
    '\n'
    'Response: list of {\n'
    '  dataElement, orgUnit, period, categoryOptionCombo,\n'
    '  value, mean, stdDev, zScore, followUp\n'
    '}'
)
sp()

h2('4.2  DTP1/DTP3 consistency')
body('Uses DHIS2\'s validation rule engine. The EPI Aggregate Metadata Package v1.1.0 ships with a built-in validation rule: Penta3 <= Penta1. The agent triggers a fresh validation run for the target period, then reads the results. Both steps are required in DHIS2 2.40 — results are not stored persistently until the run is explicitly triggered.')
sp()
body('Flagging thresholds vary by data level (from AHEAD methodology):')
tbl(
    ['Data level', 'Absolute threshold', 'Relative threshold', 'Logic'],
    [
        ['Monthly facility data',  '100 doses',   '>30% difference', 'Either condition triggers a flag'],
        ['Monthly admin2 data',    '250 doses',   '>20% difference', 'Either condition'],
        ['Annual admin2 data',     '1,000 doses', '>15% difference', 'Either condition'],
    ]
)
sp()
body('For the prototype (monthly facility data): (Penta3 - Penta1) > 100 doses OR (Penta3 - Penta1) / Penta1 > 30%. In the demo scenario (Penta1=45, Penta3=80), the relative difference is 78% — flagged on the relative condition even though the absolute gap (35) is below 100.')
mono(
    '# Step 1 — trigger validation run\n'
    'POST /api/validation\n'
    'Body: {\n'
    '  "organisationUnit": "RFhqluFmvRG",\n'
    '  "startPeriod": "202604",\n'
    '  "endPeriod":   "202604"\n'
    '}\n'
    '\n'
    '# Step 2 — read results\n'
    'GET /api/validationResults\n'
    '  ?ou=RFhqluFmvRG\n'
    '  &startPeriod=202604\n'
    '  &endPeriod=202604\n'
    '\n'
    'Response: list of {\n'
    '  validationRule {id, name},\n'
    '  orgUnit {id, name},\n'
    '  period,\n'
    '  leftsideValue,   # Penta3 value\n'
    '  rightsideValue   # Penta1 value\n'
    '}'
)
sp()

h2('4.3  Missing report detection')
body('Queries the complete dataset registration log to find which facilities submitted for the target period. Compares against the expected list (all level-5 org units assigned to the EPI dataset). Any facility present in expected but absent from registrations is flagged. Runs daily starting on day 5 of the following month.')
mono(
    'GET /api/completeDataSetRegistrations\n'
    '  ?dataSet=vI4ihClxSm4\n'
    '  &orgUnit=RFhqluFmvRG\n'
    '  &startPeriod=202604\n'
    '  &endPeriod=202604\n'
    '  &children=true\n'
    '\n'
    'Expected facilities: all org units at level 5 with dataset vI4ihClxSm4 assigned.\n'
    'Missing = expected_set - {r["organisationUnit"]["id"] for r in response}'
)
sp()

h2('4.4  Applying corrections')
body('When a facility worker confirms a corrected value via SMS, the agent writes it back to DHIS2 using the standard dataValueSets API — identical to any user submitting data — authenticated as agent_service. After writing, the agent re-runs all three checks for that facility-period to confirm the correction resolved the issue.')
mono(
    'POST /api/dataValueSets\n'
    'Body: {\n'
    '  "dataSet":    "vI4ihClxSm4",\n'
    '  "orgUnit":    "<facility UID>",\n'
    '  "period":     "202604",\n'
    '  "dataValues": [\n'
    '    {\n'
    '      "dataElement":        "<DE UID>",\n'
    '      "categoryOptionCombo": "<COC UID>",\n'
    '      "value":              "35"\n'
    '    }\n'
    '  ]\n'
    '}'
)
sp()

h2('4.5  Response option schema (from AHEAD methodology)')
body('Every SMS notification presents numbered multiple-choice options — the same fixed response categories in the current AHEAD Excel dropdowns. These are defined per check type and never change between conversations. They are not generated by Claude; Claude only parses which number was selected and extracts any required follow-up value.')
sp()
tbl(
    ['Check type', 'Option', 'Label', 'Requires follow-up?', 'Agent action'],
    [
        ['Outlier', '1', 'Replace with 6-month avg',    'No',  'Agent calculates avg of 3mo before + 3mo after; applies immediately'],
        ['Outlier', '2', 'Keep as-is (explain)',         'Yes — short explanation text', 'Log value unchanged with comment'],
        ['Outlier', '3', 'Treat as zero',                'No',  'Set value to 0'],
        ['Outlier', '4', 'Replace with specific number', 'Yes — numeric value', 'Apply provided value'],
        ['Outlier', '5', 'Use facility doses only',      'No',  'Pull from HF column; overwrite total'],
        ['Outlier', '6', 'Use outreach doses only',      'No',  'Pull from outreach column; overwrite total'],
        ['DTP inconsistency', '1', 'Keep as-is (explain)',        'Yes — explanation', 'Log unchanged with comment'],
        ['DTP inconsistency', '2', 'Set both to Penta1 value',    'No',  'Write Penta1 value to both Penta1 and Penta3'],
        ['DTP inconsistency', '3', 'Set both to Penta3 value',    'No',  'Write Penta3 value to both'],
        ['DTP inconsistency', '4', 'Replace with specific values', 'Yes — Penta1,Penta3 separated by comma', 'Apply both values'],
        ['DTP inconsistency', '5', 'Other (explain)',              'Yes — explanation', 'Log comment; no data change'],
        ['Missing report', 'SUBMIT', 'Submitting now',             'No (re-check in 1h)',  'Schedule re-check; close auto if found'],
        ['Missing report', '1', 'Replace with 6-month avg',        'No',  'Impute avg for all vaccines'],
        ['Missing report', '2', 'Replace with zero',               'No',  'Write 0 for all vaccines (closed/no service)'],
        ['Missing report', '3', 'Replace with specific values',     'Yes — BCG,P1,P3,MR1 comma-separated', 'Apply per-vaccine values'],
        ['Missing report', '4', 'Remove from analysis',            'No',  'Log as excluded; no data written'],
    ]
)
sp()

h2('4.7  Deduplication')
body('Before creating a new issue, the agent queries the database for any open issue with the same (org_unit_uid, period, check_type, data_element_uid). If one exists and is not yet resolved or dismissed, the detection is skipped silently. This prevents duplicate SMS notifications if the same facility submits an update and the violation is still present on the next poll cycle.')
sp()

# ── 5. Database Schema ────────────────────────────────────────────────────────
h1('5. Database Schema (SQLite → PostgreSQL in production)')
body('Three tables. All timestamps stored as ISO-8601 UTC strings. The schema is PostgreSQL-compatible — migration from SQLite requires only changing the connection string and TEXT PRIMARY KEY to SERIAL/UUID where appropriate.')
sp()

h2('issues')
mono(
    'CREATE TABLE issues (\n'
    '  id               TEXT PRIMARY KEY,      -- DQ-AX42\n'
    '  check_type       TEXT NOT NULL,         -- outlier | missing_report | dtp_inconsistency\n'
    '  org_unit_uid     TEXT NOT NULL,\n'
    '  facility_name    TEXT NOT NULL,\n'
    '  woreda_name      TEXT NOT NULL,\n'
    '  period           TEXT NOT NULL,         -- YYYYMM e.g. 202604\n'
    '  data_element_uid TEXT,                  -- NULL for missing_report\n'
    '  data_element_name TEXT,\n'
    '  coc_uid          TEXT,                  -- categoryOptionCombo UID\n'
    '  flagged_value    TEXT,                  -- NULL for missing_report\n'
    '  selected_option  TEXT,                  -- "1"-"6", "SUBMIT" — set on first inbound reply\n'
    '  resolved_value   TEXT,                  -- final value or action applied\n'
    '  status           TEXT NOT NULL DEFAULT \'notified\',\n'
    '                                          -- notified | awaiting_followup\n'
    '                                          -- | resolved | escalated\n'
    '  cascade_level    INTEGER NOT NULL DEFAULT 1,\n'
    '                                          -- 1=facility 2=woreda 3=zone\n'
    '                                          -- 4=region   5=national\n'
    '  retry_count      INTEGER NOT NULL DEFAULT 0,\n'
    '  opened_at        TEXT NOT NULL,\n'
    '  last_contact_at  TEXT,\n'
    '  resolved_at      TEXT\n'
    ');'
)
sp()

h2('conversations')
mono(
    'CREATE TABLE conversations (\n'
    '  id        INTEGER PRIMARY KEY AUTOINCREMENT,\n'
    '  issue_id  TEXT NOT NULL REFERENCES issues(id),\n'
    '  direction TEXT NOT NULL,    -- outbound | inbound\n'
    '  phone     TEXT NOT NULL,    -- E.164 format\n'
    '  body      TEXT NOT NULL,\n'
    '  sent_at   TEXT NOT NULL     -- ISO-8601 UTC\n'
    ');'
)
sp()

h2('contacts')
body('Phone numbers (and email for Phase 2) are maintained in this table by the AHEAD team, not read from DHIS2 user profiles. George\'s team does not maintain a phone registry inside DHIS2. The email column is nullable and present from day one so adding email in Phase 2 requires no schema migration.')
mono(
    'CREATE TABLE contacts (\n'
    '  org_unit_uid  TEXT    NOT NULL,\n'
    '  level         INTEGER NOT NULL,  -- 1=national 2=region 3=zone 4=woreda 5=facility\n'
    '  name          TEXT    NOT NULL,  -- person name\n'
    '  phone         TEXT,              -- E.164 format; NULL if no SMS contact\n'
    '  email         TEXT,              -- NULL if no email contact (Phase 2)\n'
    '  PRIMARY KEY (org_unit_uid, level)\n'
    ');\n'
    '\n'
    '-- At least one of phone or email must be non-null.\n'
    '-- MVP: agent uses phone only. Email added in Phase 2 with no schema change.\n'
    '-- Maintained by the AHEAD team; seeded from a CSV for the prototype.'
)
sp()

# ── 6. Cascade State Machine ──────────────────────────────────────────────────
h1('6. Cascade State Machine')
mono(
    '                   Issue detected\n'
    '                        │\n'
    '                        ▼\n'
    '               ┌─────────────────┐\n'
    '               │    NOTIFIED     │  Initial SMS sent to facility\n'
    '               └────────┬────────┘\n'
    '                        │\n'
    '           ┌────────────┼──────────────────┐\n'
    '    Reply  │            │ No reply 24h      │ KEEP\n'
    '  (value)  │            │ (retry_count < 3) │\n'
    '           │            ▼                   ▼\n'
    '           │     retry SMS sent      ┌─────────────┐\n'
    '           │     retry_count++       │  DISMISSED  │\n'
    '           │            │            └─────────────┘\n'
    '           │    3 retries exhausted\n'
    '           │            │\n'
    '           │            ▼\n'
    '           │    ┌────────────────┐\n'
    '           │    │   ESCALATED    │  Summary sent to next level\n'
    '           │    │ cascade_level++│  (woreda → zone → region → national)\n'
    '           │    └────────────────┘\n'
    '           │\n'
    '           ▼\n'
    '   ┌────────────────────┐\n'
    '   │  AWAITING_CONFIRM  │  Agent proposes specific change\n'
    '   └──────┬─────────────┘\n'
    '          │\n'
    '    ┌─────┼────────────────────┐\n'
    '    YES   NO / different value  │\n'
    '    │     │                    │\n'
    '    ▼     ▼                    │\n'
    '┌──────┐  Re-send confirmation  │\n'
    '│RESOL-│  with updated value    │\n'
    '│ VED  │  (loop back up)        │\n'
    '└──────┘'
)
sp()
tbl(
    ['Status', 'Trigger', 'Agent action'],
    [
        ['notified',
         'Issue first detected by DQ engine',
         'Send numbered-option SMS to facility contact; record last_contact_at'],
        ['notified → retry',
         'No inbound reply within 24h; retry_count < 3',
         'Re-send abbreviated option-list SMS; increment retry_count'],
        ['notified → escalated',
         'retry_count reaches 3 (72h elapsed)',
         'Send summary SMS to next-level contact; increment cascade_level'],
        ['notified → awaiting_followup',
         'Inbound option number that requires additional input (e.g. option 4 = specific number)',
         'Claude extracts option; agent sends follow-up question; record selected_option'],
        ['notified → resolved',
         'Inbound option number that needs no additional input (e.g. option 1 = 6-month avg)',
         'Apply action immediately (compute avg / write zero / etc.); write to DHIS2; send resolution SMS; close'],
        ['awaiting_followup → resolved',
         'Inbound follow-up value (specific number, explanation text, comma-separated values)',
         'Claude extracts value; apply to DHIS2 if needed; send resolution SMS; close issue'],
        ['notified → awaiting_followup (SUBMIT)',
         'Inbound SUBMIT on missing report',
         'Schedule 1-hour re-check; send acknowledgement SMS'],
        ['awaiting_followup (SUBMIT) → resolved',
         '1-hour re-check finds the registration',
         'Auto-close; send confirmation SMS'],
    ]
)
sp()
h2('Escalation timing')
tbl(
    ['Level', 'Recipient', 'Trigger', 'SLA to next level'],
    [
        ['1 — Facility',  'Facility worker',         'Issue detected',                '72h (3 × 24h retries)'],
        ['2 — Woreda',    'Woreda HMIS officer',      '72h no facility resolution',    '1 week'],
        ['3 — Zone',      'Zone officer',             '1 week no woreda resolution',   '1 week'],
        ['4 — Region',    'Regional focal point',     '1 week no zone resolution',     'Until end of month'],
        ['5 — National',  'National TWG',             'End of month digest (all open issues)', 'N/A'],
    ]
)
sp()

# ── 7. Claude API Integration ─────────────────────────────────────────────────
h1('7. Claude API Integration')
body('Claude is used for two narrow tasks. The response options themselves are fixed (see Section 4.5) — Claude does not generate or choose them. Its job is: (1) determine which numbered option the worker selected and extract any required follow-up value; (2) generate the variable parts of outbound messages (facility name, flagged value, normal range). Hallucination risk is very low — both tasks are tightly scoped extractions, not open-ended generation.')
sp()
tbl(
    ['Task', 'Model', 'Input', 'Output', 'Max tokens'],
    [
        ['Parse inbound reply',
         'claude-haiku-4-5',
         'SMS body + issue context (check type, option list, current state)',
         'JSON: {"option": <1-6|"SUBMIT"|null>, "followup_value": <string>|null, "state": "option_selected"|"followup_provided"|"unknown"}',
         '80'],
        ['Generate outbound message body',
         'claude-haiku-4-5',
         'Message type (notification|followup_request|resolution) + issue context (facility, vaccine, flagged value, normal range)',
         'SMS text filling in the variable parts of the template. Fixed option list is appended by the agent, not Claude.',
         '100'],
    ]
)
sp()
body('Prompt caching (cache_control: ephemeral) is applied to the system prompt for each task. The user turn (specific SMS body or issue values) is not cached.')
sp()
mono(
    '# Parse inbound — system prompt (cached)\n'
    '"The worker is responding to an outlier flag. Options presented were:\n'
    ' 1=6mo-avg 2=keep(explain) 3=zero 4=specific-number 5=HF-only 6=outreach-only\n'
    ' Current state: awaiting_option.\n'
    ' Extract the selected option number and any follow-up value from the reply.\n'
    ' Respond with JSON only: {option, followup_value, state}\n'
    ' state: option_selected (just a number), followup_provided (value after option 4 requested), unknown"\n'
    '\n'
    '# User turn: worker replies "4"\n'
    '{"option": 4, "followup_value": null, "state": "option_selected"}\n'
    '→ agent sends follow-up: "Reply with the correct BCG under-1 value."\n'
    '\n'
    '# User turn: worker replies "35" (after follow-up)\n'
    '{"option": 4, "followup_value": "35", "state": "followup_provided"}\n'
    '→ agent applies correction to DHIS2, sends resolution SMS'
)
sp()

# ── 8. SMS Integration ────────────────────────────────────────────────────────
h1('8. SMS Integration (Twilio)')
h2('Outbound')
body('Agent calls the Twilio Messages API directly via HTTPS. One SMS per issue event. Woreda-level and above notifications are batched summaries (one SMS per summary, not one per issue).')
mono(
    'POST https://api.twilio.com/2010-04-01/Accounts/{TWILIO_ACCOUNT_SID}/Messages\n'
    '  From = TWILIO_FROM_NUMBER   (E.164 format, e.g. +12015551234)\n'
    '  To   = <recipient phone from contacts table>\n'
    '  Body = <message text>\n'
    '\n'
    '  On success: HTTP 201, sid returned and logged to conversations table\n'
    '  On failure: log error, retry once after 60s, then mark issue with send_error flag'
)
sp()

h2('Inbound (webhook routing)')
body('Twilio posts inbound SMS to POST /webhook/sms. The agent extracts the reference ID from the message body using a regex, looks up the corresponding open issue, and routes the body to the reply handler.')
mono(
    'POST /webhook/sms\n'
    '  Twilio params: From (sender phone), Body (SMS text), MessageSid\n'
    '\n'
    '  1. Extract ref ID:  re.search(r\'DQ-[A-Z0-9]{4}\', body.upper())\n'
    '  2. Query DB:        SELECT * FROM issues WHERE id=? AND status NOT IN (\'resolved\',\'dismissed\')\n'
    '  3. Not found:       reply "Ref ID not recognised. View issues at <url>"\n'
    '  4. Found:           pass to reply_handler(issue, sms_body)\n'
    '                         → Claude parse\n'
    '                         → state transition\n'
    '                         → Twilio response (TwiML or REST post-reply)\n'
    '  5. No ref ID at all: check if From number has exactly one open issue\n'
    '                         → if yes, route to that issue\n'
    '                         → if multiple, ask worker to include the ref ID'
)
sp()

h2('Reference ID format')
body('Format: DQ-XXXX (4-character suffix). Character set excludes visually ambiguous characters (0/O, 1/I/L, 2/Z). Remaining set: A B C D E F G H J K M N P Q R S T U V W X Y  +  3 4 5 6 7 8 9 = 29 characters. DQ-XXXX gives 29^4 = 707,281 unique IDs — sufficient for MVP. IDs are never reused; resolved issues keep their original ID for audit purposes.')
sp()

# ── 9. Flask Application ──────────────────────────────────────────────────────
h1('9. Flask Application (Standalone Web App)')
body('A single Flask process on port 5001 handles the inbound SMS webhook and serves the issue log web app. The issue log is a standalone web application — completely separate from DHIS2, no DHIS2 login required, readable by anyone with network access to port 5001. It reads from the agent\'s own SQLite database. Runs in the same Docker container as the scheduler and DQ engine.')
sp()
tbl(
    ['Endpoint', 'Method', 'Description'],
    [
        ['/',              'GET',  'Redirect to /issues'],
        ['/issues',        'GET',  'Issue log: HTML table of all issues, filterable by status / woreda / period. Conversation thread expandable per row.'],
        ['/issues/<id>',   'GET',  'Detail view: full SMS thread for one issue, DHIS2 data element values, timeline of state changes.'],
        ['/webhook/sms',   'POST', 'Twilio inbound SMS webhook. Validates Twilio signature, routes reply to state machine.'],
        ['/api/status',    'GET',  'JSON health check: uptime, open issue count, last successful poll timestamp, last error.'],
        ['/api/scan',      'POST', 'Trigger an immediate DQ check without waiting for the next 30-second poll cycle. Runs outlier + DTP checks across all facilities (ignores lastUpdated scoping). Optional ?check=missing_reports to run only the completeness check. Returns JSON of new issues created. Used for demos and manual testing.'],
    ]
)
sp()

# ── 10. APScheduler Jobs ──────────────────────────────────────────────────────
h1('10. APScheduler Jobs')
body('APScheduler (BackgroundScheduler) runs inside the Flask process. All times East Africa Time (UTC+3). Jobs are idempotent — restarting the container does not create duplicate issues or duplicate SMS messages.')
sp()
tbl(
    ['Job ID', 'Schedule', 'What it does'],
    [
        ['poll_changes',
         'Every 30 seconds',
         'Queries GET /api/dataValueSets?lastUpdated={last_checked}. If the response is empty, does nothing. If non-empty, extracts the changed (orgUnit, period) pairs and passes them to run_dq_checks as a targeted scope. Updates last_checked timestamp on every cycle.'],
        ['run_dq_checks',
         'On-demand (triggered by poll_changes when changes detected)',
         'Runs outlier detection and DTP validation only for the (orgUnit, period) pairs identified by poll_changes. Not on a fixed schedule — fires only when new data is present. Deduplicates against open issues before creating new records.'],
        ['check_missing_reports',
         'Daily at 08:00 EAT',
         'Queries completeDataSetRegistrations for prior month. Flags facilities with no registration. Only runs if today >= MISSING_REPORT_START_DAY (default: 10th of the month). Configurable so programs can match their own reporting deadline.'],
        ['process_timers',
         'Every 30 minutes',
         'Scans open issues: sends retry SMS if 24h elapsed and retry_count < 3; escalates to next cascade level if 72h elapsed at current level; escalates zone→region→national on weekly cadence.'],
        ['monthly_summary',
         '1st of month at 09:00 EAT',
         'Sends digest of all unresolved issues from prior month to national TWG contact. Includes count by woreda and issue type.'],
    ]
)
sp()

# ── 11. Docker Compose Deployment ─────────────────────────────────────────────
h1('11. Docker Compose Deployment')
body('The agent is added as a fourth service to docker-compose.yml alongside the existing db, web (DHIS2), and any future services. Both the agent and DHIS2 containers share the same Docker bridge network; the agent reaches DHIS2 at http://web:8080/api (internal DNS), not localhost.')
sp()
mono(
    '# Addition to docker-compose.yml\n'
    '  agent:\n'
    '    build: ./agent              # Dockerfile at ./agent/Dockerfile\n'
    '    ports:\n'
    '      - "5001:5001"\n'
    '    volumes:\n'
    '      - ./agent/db:/data        # SQLite file persisted to host filesystem\n'
    '      - ./agent/contacts.csv:/app/contacts.csv  # contact registry (read-only)\n'
    '    environment:\n'
    '      DHIS2_BASE_URL:       http://web:8080/api\n'
    '      AGENT_USER:           ${AGENT_USER}\n'
    '      AGENT_PASS:           ${AGENT_PASS}\n'
    '      CLAUDE_API_KEY:       ${CLAUDE_API_KEY}\n'
    '      TWILIO_ACCOUNT_SID:   ${TWILIO_ACCOUNT_SID}\n'
    '      TWILIO_AUTH_TOKEN:    ${TWILIO_AUTH_TOKEN}\n'
    '      TWILIO_FROM_NUMBER:   ${TWILIO_FROM_NUMBER}\n'
    '      # Thresholds and timers live in config.py, not here\n'
    '    depends_on:\n'
    '      - web\n'
    '    restart: unless-stopped\n'
    '\n'
    '# Agent Dockerfile (./agent/Dockerfile)\n'
    '  FROM python:3.12-slim\n'
    '  WORKDIR /app\n'
    '  COPY requirements.txt .\n'
    '  RUN pip install -r requirements.txt\n'
    '  COPY . .\n'
    '  CMD ["python", "agent.py"]'
)
sp()

# ── 12. Configuration: Two-Layer Approach ────────────────────────────────────
h1('12. Configuration: Two-Layer Approach')
body('Configuration is split into two files by type: secrets in .env (never committed), and instance metadata in config.py (safe to commit). This separation means deploying to a new country or DHIS2 instance requires only updating config.py — no secrets need to change unless you are also changing your API provider accounts.')
sp()
mono(
    '.env          ← secrets and runtime URLs (never commit)\n'
    'config.py     ← instance UIDs, thresholds, hierarchy settings (safe to commit)\n'
    '\n'
    'For a new country deployment:\n'
    '  1. Update config.py with the country\'s DHIS2 UIDs\n'
    '  2. Set DHIS2_BASE_URL in .env to the country\'s DHIS2 URL\n'
    '  3. Set AGENT_USER / AGENT_PASS for the agent_service account on that instance\n'
    '  4. Seed the contacts table with that country\'s phone numbers\n'
    '  That\'s it — no code changes required.'
)
sp()

h2('12a.  Secrets (.env)')
body('Passwords, API keys, and deployment URLs. Never committed to version control. Copy .env.example to .env and fill in values before running anything.')
sp()
tbl(
    ['Variable', 'Used by', 'Description'],
    [
        ['DHIS2_BASE_URL',       'Setup scripts, Agent', 'http://localhost:8080/api (local) or http://web:8080/api (Docker internal)'],
        ['DHIS2_ADMIN_USER',     'Setup scripts only',   'DHIS2 admin username (demo setup only)'],
        ['DHIS2_ADMIN_PASS',     'Setup scripts only',   'DHIS2 admin password'],
        ['DHIS2_USER_PASSWORD',  'Setup scripts only',   'Shared password for demo role-based user accounts'],
        ['AGENT_USER',           'Agent',                'DHIS2 agent_service username'],
        ['AGENT_PASS',           'Agent',                'DHIS2 agent_service password'],
        ['CLAUDE_API_KEY',       'Agent',                'Anthropic API key'],
        ['TWILIO_ACCOUNT_SID',   'Agent',                'Twilio account SID'],
        ['TWILIO_AUTH_TOKEN',    'Agent',                'Twilio auth token'],
        ['TWILIO_FROM_NUMBER',   'Agent',                'Twilio sending phone number (E.164 format)'],
        ['POSTGRES_DB',          'docker-compose.yml',   'PostgreSQL database name (DHIS2 internal)'],
        ['POSTGRES_USER',        'docker-compose.yml',   'PostgreSQL username'],
        ['POSTGRES_PASSWORD',    'docker-compose.yml',   'PostgreSQL password'],
    ]
)
sp()

h2('12b.  Instance metadata (config.py)')
body('DHIS2 UIDs, DQ thresholds, hierarchy settings, and timers. Country-specific but not secret — safe to commit. Copy config.example.py to config.py. For a new country deployment, update the UIDs by querying the target DHIS2 instance\'s API; everything else can stay at its default until the country team requests tuning.')
sp()
tbl(
    ['Parameter', 'Default (Ethiopia)', 'Description'],
    [
        ['ROOT_ORG_UNIT_UID',   'RFhqluFmvRG',  'Top-level org unit the agent scans. Find via GET /api/organisationUnits?level=1'],
        ['FACILITY_LEVEL',      '5',             'Level number of facilities. Ethiopia is 5-level; a 4-level country uses 4.'],
        ['ROUTINE_DATASET_UID', 'vI4ihClxSm4',  'EPI routine vaccine delivery dataset UID. Find via GET /api/dataSets?filter=name:like:EPI'],
        ['DATA_ELEMENTS',       '(dict)',        'UIDs for BCG, Penta1, Penta3, MR1. Find via GET /api/dataElements?filter=name:like:BCG'],
        ['CATEGORY_OPTION_COMBOS', '(dict)',     'UIDs for under-1 and >= 1 year age disaggregations'],
        ['OUTLIER_Z_THRESHOLD', '3.0',           'Z-score threshold for outlier detection'],
        ['OUTLIER_ABS_THRESHOLD','100',           'Absolute dose difference from mean (AHEAD method 5 proxy)'],
        ['OUTLIER_MIN_HISTORY_MONTHS', '3',      'Minimum months of history before outlier detection runs for a facility'],
        ['DTP_THRESHOLDS',      '(dict)',        'Relative and absolute thresholds by data level (facility/admin2, monthly/annual)'],
        ['MISSING_REPORT_START_DAY', '10',       'Day of following month to begin missing report checks. Confirm deadline with country team.'],
        ['RETRY_INTERVAL_HOURS', '24',           'Hours between retries at facility level'],
        ['MAX_RETRIES',         '3',             'Retries before escalation (72h total at default)'],
        ['ESCALATION_DAYS',     '(dict)',        'Days at each level before escalating: woreda=3, zone=10, region=17'],
        ['POLL_INTERVAL_SEC',   '30',            'lastUpdated poll frequency. Controls detection latency, not API cost.'],
    ]
)
sp()

# ── 13. End-to-End Walkthrough ────────────────────────────────────────────────
h1('13. End-to-End Walkthrough: Outlier Scenario')
body('eth_facility_01 (Almaz Tadesse, Addi Arekay Health Center) submits June 2026 EPI data with BCG under-1 = 970. Correct value is 97 — an extra zero. 22-month baseline mean is ~97. The following traces the full agent lifecycle.')
sp()
mono(
    'T+0:00  eth_facility_01 submits June 2026 form in DHIS2\n'
    '         Data lands in datavalue table immediately\n'
    '\n'
    'T+0:30  poll_changes fires\n'
    '  → GET /api/dataValueSets?...&lastUpdated={T+0:00} → Addi Arekay HC / 202606 returned\n'
    '  → run_dq_checks triggered for (aV3ume00zx5, 202606)  ← Addi Arekay HC UID\n'
    '  → GET /api/outlierDetection?ou=aV3ume00zx5&... → BCG <1yr, value=970, z=~87\n'
    '  → dedup check: no open issue for (aV3ume00zx5, 202606, outlier, BCG)\n'
    '  → INSERT INTO issues (id=\'DQ-AX42\', status=\'notified\', ...)\n'
    '  → SELECT phone FROM contacts WHERE org_unit_uid=\'aV3ume00zx5\' AND level=5\n'
    '  → POST Twilio → numbered-option SMS sent to facility worker\'s phone\n'
    '  → INSERT INTO conversations (direction=\'outbound\', body=\'[DQ-AX42] ...\')\n'
    '\n'
    'T+0:31  Almaz receives SMS with 6 numbered options (see UX doc Section 3.1)\n'
    '\n'
    'T+1:00  Almaz replies: "4"\n'
    '  → Twilio → POST /webhook/sms (From=+251..., Body="4")\n'
    '  → regex finds DQ-AX42 (or matches by phone to single open issue)\n'
    '  → Claude parse: {"option": 4, "followup_value": null, "state": "option_selected"}\n'
    '  → UPDATE issues SET status=\'awaiting_followup\', selected_option=\'4\'\n'
    '  → Send follow-up: "[DQ-AX42] Option 4: replace with specific number. Reply with correct BCG under-1."\n'
    '\n'
    'T+1:01  Almaz replies: "97"\n'
    '  → Claude parse: {"option": 4, "followup_value": "97", "state": "followup_provided"}\n'
    '  → POST /api/dataValueSets (as agent_service, value=97) → DHIS2 returns 200\n'
    '  → UPDATE issues SET status=\'resolved\', resolved_value=\'97\', resolved_at=now()\n'
    '  → Send resolution: "[DQ-AX42] Done. BCG under-1 updated to 97. Decision logged. Issue closed."\n'
    '  → Re-run outlier check for Addi Arekay HC / 202606 → no violations (z-score now ~0)\n'
    '\n'
    'T+1:02  Issue DQ-AX42 visible in /issues page as Resolved'
)
sp()

doc.save('AHEAD_AI_Tech_Architecture.docx')
print('Saved: AHEAD_AI_Tech_Architecture.docx')
